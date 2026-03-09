"""
╔══════════════════════════════════════════════════════════════╗
║           core_scraper.py — Backend Engine                   ║
║                                                              ║
║  Універсальний парсер з Readability-підходом                ║
║  Selenium (JS disabled) → trafilatura → readability-lxml    ║
║  → власна евристика → DOCX/PDF                             ║
╚══════════════════════════════════════════════════════════════╝
"""

# ═══════════════════════════════════════════════════════════════
#                        ІМПОРТИ
# ═══════════════════════════════════════════════════════════════

import os
import re
import time
import json
import platform
import subprocess
import hashlib
import warnings
import urllib.parse
from io import BytesIO
from datetime import datetime
from copy import deepcopy

from PyQt6.QtCore import QThread, pyqtSignal

# ── Selenium ──
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# ── Парсинг та обробка ──
from bs4 import BeautifulSoup, Tag, NavigableString
from PIL import Image
import requests

# ── Readability-бібліотеки (Universal Content Extraction) ──
try:
    import trafilatura
    from trafilatura.settings import use_config as trafilatura_use_config

    HAS_TRAFILATURA = True
except ImportError:
    HAS_TRAFILATURA = False

try:
    from readability import Document as ReadabilityDocument

    HAS_READABILITY = True
except ImportError:
    HAS_READABILITY = False

# ── Документи ──
from docx import Document as WordDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Переклад ──
from deep_translator import MicrosoftTranslator, GoogleTranslator

# ── PDF конвертація ──
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

warnings.filterwarnings("ignore", category=UserWarning, module='requests')

# ── Імпорт теми (для кольорів у сигналах) ──
from config import Theme

# ═══════════════════════════════════════════════════════════════
#            КОНСТАНТИ ДЛЯ УНІВЕРСАЛЬНОГО ПАРСИНГУ
# ═══════════════════════════════════════════════════════════════

# Теги, які ЗАВЖДИ видаляємо
STRIP_TAGS = frozenset([
    'nav', 'aside', 'footer', 'header', 'script', 'style', 'noscript',
    'button', 'form', 'svg', 'input', 'select', 'textarea', 'label',
    'dialog', 'menu', 'menuitem', 'template',
])

# Ключові слова у class/id, що вказують на сміття
JUNK_PATTERNS = re.compile(
    r'(?:newsletter|promo(?:tion)?|recirc|related[-_]?(?:articles?|posts?|stories)?|'
    r'recommend(?:ed|ation)?|social[-_]?(?:share|media|links?|buttons?)|'
    r'share[-_]?(?:bar|buttons?|widget)|author[-_]?bio|sidebar|bottom[-_]?bar|'
    r'ad[-_](?:slot|container|wrapper|unit|banner)|advertisement|trending|'
    r'popular[-_]?(?:articles?|posts?)|comment(?:s|[-_]section)?|disqus|'
    r'subscribe|signup|sign[-_]up|follow[-_]us|cookie[-_]?(?:banner|consent|notice)|'
    r'gdpr|privacy[-_]?(?:banner|notice)|popup|modal|overlay|'
    r'breadcrumb|pagination|pager|navigation|nav[-_]|menu[-_]|'
    r'taboola|outbrain|mgid|teads|recirculation|'
    r'read[-_]?(?:more|next|also)|more[-_]?(?:from|stories|articles)|'
    r'most[-_]?(?:read|popular|viewed)|top[-_]?stories|'
    r'widget[-_]?(?:area|container)|'
    r'footer[-_]|header[-_]|masthead|site[-_]?header|'
    r'skip[-_]?(?:to|link)|accessibility|screen[-_]?reader|'
    r'print[-_]?(?:only|hide)|no[-_]?print|'
    r'embed[-_]?(?:social|code)|infogram|datawrapper)',
    re.IGNORECASE
)

# Ключові слова у class/id, що вказують на КОНТЕНТ
CONTENT_PATTERNS = re.compile(
    r'(?:article[-_]?(?:body|content|text|main|full)|'
    r'story[-_]?(?:body|content|text)|'
    r'post[-_]?(?:body|content|text|entry)|'
    r'entry[-_]?(?:content|body|text)|'
    r'main[-_]?(?:content|text|body|article)|'
    r'body[-_]?(?:content|text|copy)|'
    r'content[-_]?(?:body|area|main|article|well)|'
    r'text[-_]?(?:body|content|block)|'
    r'single[-_]?(?:content|post|article)|'
    r'news[-_]?(?:body|content|text|article)|'
    r'blog[-_]?(?:body|content|text|post)|'
    r'rich[-_]?text|prose|markdown[-_]?body)',
    re.IGNORECASE
)

# Спам-фрази, які вказують на нерелевантний текст
SPAM_PHRASES = [
    "підпишіться на", "subscribe to", "sign up for",
    "read more:", "more from", "newsletter", "follow us",
    "share this", "поділитися", "recommended for you",
    "you may also like", "don't miss", "trending now",
    "click here", "натисніть тут", "завантажити додаток",
    "download the app", "install our", "join our",
    "get unlimited", "start your free trial",
    "already a subscriber", "create an account",
    "log in to", "увійдіть щоб", "зареєструйтесь",
    "cookies", "privacy policy", "terms of service",
    "all rights reserved", "всі права захищені",
    "copyright ©", "associated press", "getty images",
]

# Мінімальна довжина "корисного" тексту
MIN_TEXT_LENGTH = 25
MIN_PARAGRAPH_WORDS = 4

# Мінімальна кількість слів для "знайдено контент"
MIN_CONTENT_WORDS = 50


# ═══════════════════════════════════════════════════════════════
#                 SCRAPING WORKER (QThread)
# ═══════════════════════════════════════════════════════════════

class ScrapingWorker(QThread):
    """
    Воркер для парсингу статей у фоновому потоці.
    Використовує каскадний підхід:
      1. trafilatura (найкращий для статей)
      2. readability-lxml (Mozilla Readability)
      3. Власна евристика text-density ratio

    Медіа (img, video, iframe) витягуються окремим пасом
    з оригінального DOM для збереження хронологічного порядку.

    JavaScript ВИМКНЕНО в Selenium (як вимога замовника).
    """

    # ── Сигнали ──
    progress_updated = pyqtSignal(int)
    status_updated = pyqtSignal(str, str)
    finished_success = pyqtSignal(str)
    finished_all = pyqtSignal()
    error_occurred = pyqtSignal(str)
    notification_requested = pyqtSignal(str, str)

    def __init__(self, urls, state, locales, parent=None):
        super().__init__(parent)
        self.urls = urls
        self.state = state.copy()
        self.locales = locales
        self._cancelled = False

        self.http_headers = {
            'User-Agent': (
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                '(KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'
            ),
            'Accept': 'image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9,uk;q=0.8',
        }

        self.translation_languages = {
            "Українська": "uk", "English": "en", "Polski": "pl",
            "Deutsch": "de", "Français": "fr", "Español": "es"
        }

        self.cache_dir = "backup_cache"
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)

    def cancel(self):
        """Скасувати обробку"""
        self._cancelled = True

    def t(self, key):
        """Отримати переклад ключа локалізації"""
        lang = self.state.get("ui_language", "uk")
        return self.locales.get(lang, self.locales["uk"]).get(key, key)

    # ═══════════════════════════════════════════════════════════
    #                    БРАУЗЕР (Selenium)
    # ═══════════════════════════════════════════════════════════

    def get_browser_driver(self):
        """
        Запускає headless браузер з ВИМКНЕНИМ JavaScript.
        Спочатку пробує Edge, потім Chrome.
        """
        # Вимикаємо JavaScript через preferences
        prefs = {
            "profile.managed_default_content_settings.javascript": 2
        }

        common_args = [
            "--headless=new",
            "--disable-gpu",
            "--no-sandbox",
            "--disable-dev-shm-usage",
            "--disable-extensions",
            "--disable-plugins",
            "--disable-infobars",
            "--disable-notifications",
            "--disable-popup-blocking",
            "--window-size=1920,1080",
            "--ignore-certificate-errors",
        ]

        # Спроба 1: Edge
        try:
            options = EdgeOptions()
            for arg in common_args:
                options.add_argument(arg)
            options.add_experimental_option("prefs", prefs)
            options.add_argument(
                f"user-agent={self.http_headers['User-Agent']}"
            )
            return webdriver.Edge(options=options)
        except Exception:
            pass

        # Спроба 2: Chrome
        try:
            options = ChromeOptions()
            for arg in common_args:
                options.add_argument(arg)
            options.add_experimental_option("prefs", prefs)
            options.add_argument(
                f"user-agent={self.http_headers['User-Agent']}"
            )
            return webdriver.Chrome(options=options)
        except Exception:
            raise Exception(
                "Не вдалося запустити ні Edge, ні Chrome браузер на цьому ПК."
            )

    # ═══════════════════════════════════════════════════════════
    #                    ПЕРЕКЛАД
    # ═══════════════════════════════════════════════════════════

    def translate_text(self, text):
        """Перекладає текст обраним рушієм з fallback"""
        if not text or len(text.strip()) < 5:
            return text

        target_name = self.state.get("target_lang_name", "Українська")
        target_code = self.translation_languages.get(target_name, "uk")
        engine = self.state.get("translation_engine", "Google Translator")

        # Розбиваємо довгий текст на чанки (ліміт API ~ 5000 символів)
        max_chunk = 4500
        if len(text) <= max_chunk:
            return self._translate_chunk(text, target_code, engine)

        chunks = self._split_text_into_chunks(text, max_chunk)
        translated_chunks = []
        for chunk in chunks:
            if self._cancelled:
                return text
            translated_chunks.append(
                self._translate_chunk(chunk, target_code, engine)
            )
        return " ".join(translated_chunks)

    def _split_text_into_chunks(self, text, max_size):
        """Розбиває текст на чанки по реченнях"""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current = ""
        for sent in sentences:
            if len(current) + len(sent) + 1 > max_size:
                if current:
                    chunks.append(current.strip())
                current = sent
            else:
                current = f"{current} {sent}" if current else sent
        if current:
            chunks.append(current.strip())
        return chunks if chunks else [text]

    def _translate_chunk(self, text, target_code, engine):
        """Перекладає один чанк тексту"""
        if engine == "Microsoft Translator":
            try:
                return MicrosoftTranslator(target=target_code).translate(text)
            except Exception:
                try:
                    return GoogleTranslator(
                        source='auto', target=target_code
                    ).translate(text)
                except Exception:
                    return text
        else:
            try:
                return GoogleTranslator(
                    source='auto', target=target_code
                ).translate(text)
            except Exception:
                try:
                    return MicrosoftTranslator(
                        target=target_code
                    ).translate(text)
                except Exception:
                    return text

    # ═══════════════════════════════════════════════════════════
    #              УТИЛІТИ ЗОБРАЖЕНЬ / URL
    # ═══════════════════════════════════════════════════════════

    def convert_image_for_docx(self, image_bytes):
        """Конвертує зображення у JPEG для вставки в DOCX"""
        try:
            img = Image.open(BytesIO(image_bytes))

            # Перевіряємо мінімальний розмір (відсіюємо іконки/піксель-трекери)
            if img.width < 80 or img.height < 80:
                return None

            # Якщо це анімований GIF — беремо перший кадр
            if hasattr(img, 'n_frames') and img.n_frames > 1:
                img.seek(0)

            if img.mode in ("RGBA", "P", "LA", "PA"):
                background = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                background.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
                img = background
            elif img.mode != "RGB":
                img = img.convert("RGB")

            output = BytesIO()
            img.save(output, format="JPEG", quality=85, optimize=True)
            output.seek(0)
            return output
        except Exception:
            return None

    def resolve_url(self, url, base_url):
        """Перетворює відносний URL на абсолютний"""
        if not url:
            return ""
        url = url.strip()
        if url.startswith('data:'):
            return ""
        if url.startswith('http'):
            return url
        if url.startswith('//'):
            return 'https:' + url
        return urllib.parse.urljoin(base_url, url)

    def _get_best_image_url(self, img_tag, base_url):
        """
        Витягує найкращий URL зображення з тегу <img>,
        враховуючи srcset, data-src, lazy-loading тощо.
        """
        candidates = []

        # Пріоритети джерел
        for attr in ['src', 'data-src', 'data-lazy-src', 'data-original',
                     'data-full-src', 'data-hi-res-src']:
            val = img_tag.get(attr)
            if val:
                resolved = self.resolve_url(val, base_url)
                if resolved and resolved.startswith('http'):
                    candidates.append(resolved)

        # Обробка srcset — обираємо найбільше зображення
        srcset = img_tag.get('srcset') or img_tag.get('data-srcset')
        if srcset:
            best_width = 0
            best_url = ""
            for entry in srcset.split(','):
                parts = entry.strip().split()
                if len(parts) >= 1:
                    url = self.resolve_url(parts[0], base_url)
                    width = 0
                    if len(parts) >= 2:
                        w_match = re.search(r'(\d+)w', parts[1])
                        if w_match:
                            width = int(w_match.group(1))
                    if url and url.startswith('http'):
                        if width > best_width:
                            best_width = width
                            best_url = url
            if best_url:
                candidates.insert(0, best_url)  # Пріоритет — srcset

        # Повертаємо першу валідну URL
        for url in candidates:
            if not any(skip in url.lower() for skip in [
                'pixel', 'tracker', 'beacon', 'spacer', 'blank',
                '1x1', 'transparent', 'placeholder', 'loading',
                'spinner', 'logo', 'icon', 'favicon', 'avatar',
                'badge', 'button', 'arrow', 'widget',
            ]):
                return url

        return candidates[0] if candidates else ""

    # ═══════════════════════════════════════════════════════════
    #                  МЕТАДАНІ СТАТТІ
    # ═══════════════════════════════════════════════════════════

    def extract_metadata(self, soup, base_url):
        """Витягує заголовок, опис, автора, дату та обкладинку"""
        meta = {
            "title": "",
            "subtitle": "",
            "author": "",
            "date": "",
            "cover_image": ""
        }

        # ── OG / Twitter зображення ──
        try:
            og_img = (
                    soup.find('meta', property='og:image') or
                    soup.find('meta', attrs={'name': 'twitter:image'}) or
                    soup.find('meta', attrs={'name': 'twitter:image:src'})
            )
            if og_img and og_img.get('content'):
                meta['cover_image'] = self.resolve_url(
                    og_img['content'], base_url
                )
        except Exception:
            pass

        # ── JSON-LD (структуровані дані) ──
        try:
            json_scripts = soup.find_all('script', type='application/ld+json')
            for script in json_scripts:
                if not script.string:
                    continue
                try:
                    raw = json.loads(script.string)
                    articles = self._find_json_ld_articles(raw)
                    for article in articles:
                        if not isinstance(article, dict):
                            continue

                        if not meta['title']:
                            meta['title'] = (
                                    article.get('headline') or
                                    article.get('name') or ""
                            )
                            if isinstance(meta['title'], list):
                                meta['title'] = meta['title'][0] if meta['title'] else ""

                        if not meta['subtitle']:
                            meta['subtitle'] = article.get('description', "")
                            if isinstance(meta['subtitle'], list):
                                meta['subtitle'] = meta['subtitle'][0] if meta['subtitle'] else ""

                        if not meta['date']:
                            d = article.get('datePublished') or article.get('dateCreated', '')
                            meta['date'] = str(d).split('T')[0] if d else ''

                        if not meta['author']:
                            meta['author'] = self._extract_author_from_jsonld(article)

                        if not meta['cover_image']:
                            meta['cover_image'] = self._extract_image_from_jsonld(
                                article, base_url
                            )

                except (json.JSONDecodeError, KeyError, TypeError):
                    continue
        except Exception:
            pass

        # ── Fallback: OG / meta теги ──
        try:
            if not meta['title']:
                tag = (
                        soup.find('meta', property='og:title') or
                        soup.find('meta', attrs={'name': 'title'}) or
                        soup.find('meta', attrs={'name': 'twitter:title'})
                )
                if tag and tag.get('content'):
                    meta['title'] = tag['content']

            if not meta['subtitle']:
                tag = (
                        soup.find('meta', property='og:description') or
                        soup.find('meta', attrs={'name': 'description'}) or
                        soup.find('meta', attrs={'name': 'twitter:description'})
                )
                if tag and tag.get('content'):
                    meta['subtitle'] = tag['content']

            if not meta['author']:
                tag = (
                        soup.find('meta', attrs={'name': 'author'}) or
                        soup.find('meta', property='article:author') or
                        soup.find('meta', attrs={'name': 'article:author'})
                )
                if tag and tag.get('content'):
                    meta['author'] = tag['content']

            if not meta['date']:
                tag = (
                        soup.find('meta', property='article:published_time') or
                        soup.find('meta', attrs={'name': 'date'}) or
                        soup.find('meta', attrs={'name': 'publish-date'}) or
                        soup.find('time', attrs={'datetime': True})
                )
                if tag:
                    d = tag.get('content') or tag.get('datetime', '')
                    meta['date'] = str(d).split('T')[0] if d else ''
        except Exception:
            pass

        # Очищуємо від зайвих пробілів
        for key in meta:
            if isinstance(meta[key], str):
                meta[key] = meta[key].strip()

        return meta

    def _find_json_ld_articles(self, data):
        """Рекурсивно шукає Article/NewsArticle в JSON-LD"""
        article_types = {
            'Article', 'NewsArticle', 'ReportageNewsArticle',
            'BlogPosting', 'WebPage', 'TechArticle',
            'ScholarlyArticle', 'SocialMediaPosting',
            'AnalysisNewsArticle', 'OpinionNewsArticle',
            'ReviewNewsArticle', 'Report',
        }
        results = []

        if isinstance(data, list):
            for item in data:
                results.extend(self._find_json_ld_articles(item))
        elif isinstance(data, dict):
            # Перевіряємо @graph
            if '@graph' in data:
                results.extend(self._find_json_ld_articles(data['@graph']))

            # Перевіряємо @type
            dtype = data.get('@type', '')
            if isinstance(dtype, list):
                dtype = dtype[0] if dtype else ''
            if dtype in article_types:
                results.append(data)
            elif not results and data.get('headline'):
                # Має headline — ймовірно стаття
                results.append(data)

        return results

    def _extract_author_from_jsonld(self, article):
        """Витягує автора з JSON-LD"""
        author = article.get('author')
        if isinstance(author, list):
            names = []
            for a in author:
                if isinstance(a, dict):
                    names.append(a.get('name', ''))
                elif isinstance(a, str):
                    names.append(a)
            return ", ".join(filter(None, names))
        elif isinstance(author, dict):
            return author.get('name', '')
        elif isinstance(author, str):
            return author
        return ""

    def _extract_image_from_jsonld(self, article, base_url):
        """Витягує зображення з JSON-LD"""
        img = article.get('image') or article.get('thumbnailUrl')
        if isinstance(img, list):
            if img:
                first = img[0]
                if isinstance(first, dict):
                    return self.resolve_url(first.get('url', ''), base_url)
                elif isinstance(first, str):
                    return self.resolve_url(first, base_url)
        elif isinstance(img, dict):
            return self.resolve_url(img.get('url', ''), base_url)
        elif isinstance(img, str):
            return self.resolve_url(img, base_url)
        return ""

    # ═══════════════════════════════════════════════════════════
    #         РІВЕНЬ 1: ОЧИЩЕННЯ HTML (ЗАГАЛЬНЕ)
    # ═══════════════════════════════════════════════════════════

    def _deep_clean_html(self, soup):
        """
        Глибоке очищення HTML від сміття.
        Працює на рівні DOM перед аналізом контенту.
        """
        # 1. Видаляємо небажані теги
        for tag_name in STRIP_TAGS:
            for tag in soup.find_all(tag_name):
                try:
                    tag.decompose()
                except Exception:
                    pass

        # 2. Видаляємо приховані елементи
        for tag in soup.find_all(True):
            try:
                style = tag.get('style', '').lower()
                if 'display:none' in style.replace(' ', '') or \
                        'visibility:hidden' in style.replace(' ', '') or \
                        'opacity:0' in style.replace(' ', ''):
                    tag.decompose()
                    continue

                aria_hidden = tag.get('aria-hidden', '').lower()
                if aria_hidden == 'true' and tag.name not in ['img', 'picture', 'video']:
                    tag.decompose()
                    continue

                hidden = tag.get('hidden')
                if hidden is not None and tag.name not in ['img', 'picture', 'video']:
                    tag.decompose()
                    continue
            except Exception:
                continue

        # 3. Видаляємо за class/id паттернами
        for tag in soup.find_all(['div', 'section', 'ul', 'ol', 'span',
                                  'figure', 'aside', 'p', 'table']):
            try:
                cls = tag.get('class', [])
                cls_str = ' '.join(cls).lower() if isinstance(cls, list) else str(cls).lower()
                id_str = str(tag.get('id', '')).lower()
                role_str = str(tag.get('role', '')).lower()

                combined = f"{cls_str} {id_str} {role_str}"

                if JUNK_PATTERNS.search(combined):
                    # Не видаляємо, якщо це головний контент
                    if not CONTENT_PATTERNS.search(combined):
                        tag.decompose()
            except Exception:
                continue

        # 4. Видаляємо рекламні фрейми
        for iframe in soup.find_all('iframe'):
            try:
                src = (iframe.get('src') or '').lower()
                # Залишаємо тільки відео-iframe
                video_hosts = [
                    'youtube.com', 'youtu.be', 'vimeo.com',
                    'dailymotion.com', 'rumble.com', 'bitchute.com',
                    'odysee.com', 'twitch.tv', 'facebook.com/plugins/video',
                ]
                if not any(host in src for host in video_hosts):
                    iframe.decompose()
            except Exception:
                continue

        return soup

    # ═══════════════════════════════════════════════════════════
    #       РІВЕНЬ 2: ПОШУК КОНТЕЙНЕРА СТАТТІ
    # ═══════════════════════════════════════════════════════════

    def _find_article_container(self, soup):
        """
        Каскадний пошук контейнера з основним контентом.
        Повертає BeautifulSoup об'єкт з вмістом статті.
        """
        # Стратегія 1: Шукаємо за content-специфічними класами
        container = self._find_by_content_patterns(soup)
        if container and self._has_enough_text(container):
            return container

        # Стратегія 2: Тег <article>
        articles = soup.find_all('article')
        if articles:
            # Якщо кілька <article>, обираємо найбільший
            best = max(articles, key=lambda a: len(a.get_text()))
            if self._has_enough_text(best):
                return best

        # Стратегія 3: Тег <main>
        main_tag = soup.find('main')
        if main_tag and self._has_enough_text(main_tag):
            return main_tag

        # Стратегія 4: Роль "main"
        main_role = soup.find(attrs={'role': 'main'})
        if main_role and self._has_enough_text(main_role):
            return main_role

        # Стратегія 5: Евристика text-density ratio
        best_block = self._find_by_text_density(soup)
        if best_block:
            return best_block

        # Fallback: body
        body = soup.find('body')
        return body if body else soup

    def _find_by_content_patterns(self, soup):
        """Шукає контейнер за відомими CSS-класами контенту"""
        for tag in soup.find_all(['div', 'section', 'article', 'main']):
            try:
                cls = tag.get('class', [])
                cls_str = ' '.join(cls).lower() if isinstance(cls, list) else str(cls).lower()
                id_str = str(tag.get('id', '')).lower()
                itemprop = str(tag.get('itemprop', '')).lower()

                combined = f"{cls_str} {id_str} {itemprop}"

                if CONTENT_PATTERNS.search(combined) or 'articlebody' in combined:
                    return tag
            except Exception:
                continue
        return None

    def _find_by_text_density(self, soup):
        """
        Власна евристика: знаходить блок з найвищим
        співвідношенням тексту до HTML (text-to-tag ratio).
        """
        candidates = []

        for tag in soup.find_all(['div', 'section', 'article', 'main']):
            try:
                # Пропускаємо занадто маленькі або великі (body-level) блоки
                html_str = str(tag)
                if len(html_str) < 200:
                    continue

                text = tag.get_text(strip=True)
                text_len = len(text)

                if text_len < 200:
                    continue

                # Рахуємо параграфи
                paragraphs = tag.find_all('p')
                p_count = len([
                    p for p in paragraphs
                    if len(p.get_text(strip=True)) > MIN_TEXT_LENGTH
                ])

                if p_count < 2:
                    continue

                # Text-to-HTML ratio
                html_len = len(html_str)
                ratio = text_len / html_len if html_len > 0 else 0

                # Рахуємо "посилання" — якщо їх забагато, це навігація
                links = tag.find_all('a')
                link_text = sum(len(a.get_text(strip=True)) for a in links)
                link_ratio = link_text / text_len if text_len > 0 else 1

                if link_ratio > 0.5:
                    continue  # Більше половини тексту — посилання, це не стаття

                # Скоринг
                score = (
                        text_len * 0.3 +
                        p_count * 50 +
                        ratio * 1000 -
                        link_ratio * 500
                )

                candidates.append((tag, score))
            except Exception:
                continue

        if candidates:
            candidates.sort(key=lambda x: x[1], reverse=True)
            return candidates[0][0]

        return None

    def _has_enough_text(self, container):
        """Перевіряє, чи є в контейнері достатньо тексту"""
        if not container:
            return False
        text = container.get_text(strip=True)
        word_count = len(text.split())
        return word_count >= MIN_CONTENT_WORDS

    # ═══════════════════════════════════════════════════════════
    #    РІВЕНЬ 3: TRAFILATURA + READABILITY EXTRACTION
    # ═══════════════════════════════════════════════════════════

    def _extract_with_trafilatura(self, raw_html, url):
        """
        Використовує trafilatura для екстракції тексту.
        Повертає список текстових блоків або None.
        """
        if not HAS_TRAFILATURA:
            return None

        try:
            # Налаштовуємо trafilatura
            config = trafilatura_use_config()
            config.set("DEFAULT", "MIN_OUTPUT_SIZE", "200")
            config.set("DEFAULT", "MIN_EXTRACTED_SIZE", "100")

            # Витягуємо текст з форматуванням
            result = trafilatura.extract(
                raw_html,
                url=url,
                include_comments=False,
                include_tables=True,
                include_links=False,
                include_images=False,
                include_formatting=True,
                favor_precision=True,
                config=config,
                output_format='txt',
            )

            if not result or len(result.strip()) < 100:
                return None

            # Парсимо текст у блоки
            blocks = []
            lines = result.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Визначаємо тип блоку
                if line.startswith('# '):
                    blocks.append({"type": "h2", "data": line[2:]})
                elif line.startswith('## '):
                    blocks.append({"type": "h3", "data": line[3:]})
                elif line.startswith('### '):
                    blocks.append({"type": "h4", "data": line[4:]})
                elif line.startswith('- ') or line.startswith('* '):
                    blocks.append({"type": "li", "data": line[2:]})
                elif line.startswith('> '):
                    blocks.append({"type": "blockquote", "data": line[2:]})
                elif len(line) > MIN_TEXT_LENGTH:
                    blocks.append({"type": "p", "data": line})

            return blocks if len(blocks) >= 3 else None

        except Exception:
            return None

    def _extract_with_readability(self, raw_html, url):
        """
        Використовує readability-lxml (Mozilla Readability).
        Повертає очищений HTML контейнер або None.
        """
        if not HAS_READABILITY:
            return None

        try:
            doc = ReadabilityDocument(raw_html, url=url)
            content_html = doc.summary()

            if not content_html or len(content_html) < 200:
                return None

            soup = BeautifulSoup(content_html, 'lxml')
            text = soup.get_text(strip=True)

            if len(text.split()) < MIN_CONTENT_WORDS:
                return None

            return soup

        except Exception:
            return None

    # ═══════════════════════════════════════════════════════════
    #      РІВЕНЬ 4: ПОБУДОВА content_list (текст + медіа)
    # ═══════════════════════════════════════════════════════════

    def _is_spam_text(self, text):
        """Перевіряє, чи текст є спамом/шумом"""
        if not text:
            return True
        text_lower = text.lower().strip()

        # Занадто короткий
        if len(text) < MIN_TEXT_LENGTH:
            return True

        # Замало слів
        if len(text.split()) < MIN_PARAGRAPH_WORDS:
            return True

        # Спам-фрази (тільки для коротких текстів)
        if len(text) < 200:
            if any(phrase in text_lower for phrase in SPAM_PHRASES):
                return True

        # JavaScript/CSS артефакти
        junk_keywords = ['cookie', 'javascript', 'enable js', 'browser',
                         'adblock', 'whitelist', '{', '}', 'function()',
                         'var ', 'const ', 'document.', 'window.']
        if any(kw in text_lower for kw in junk_keywords):
            return True

        return False

    def _extract_media_from_container(self, container, base_url, cover_image=""):
        """
        Витягує медіа елементи (img, video, iframe) з контейнера
        зі збереженням їх позицій відносно текстових блоків.

        Повертає dict: позиція_елемента -> [медіа_елементи]
        """
        media_map = {}
        seen_urls = set()

        if cover_image:
            seen_urls.add(cover_image)

        # Знаходимо всі текстові/медіа елементи в порядку DOM
        all_elements = container.find_all(
            ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote',
             'img', 'picture', 'figure', 'video', 'iframe']
        )

        text_index = 0
        for el in all_elements:
            try:
                if el.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                               'li', 'blockquote']:
                    text = el.get_text(strip=True)
                    if text and len(text) > MIN_TEXT_LENGTH:
                        text_index += 1

                elif el.name == 'figure':
                    # Figure може містити img + figcaption
                    img = el.find('img')
                    if img:
                        url = self._get_best_image_url(img, base_url)
                        if url and url not in seen_urls:
                            seen_urls.add(url)
                            if text_index not in media_map:
                                media_map[text_index] = []
                            media_map[text_index].append({
                                "type": "img",
                                "data": url
                            })

                elif el.name == 'picture':
                    img = el.find('img')
                    if img:
                        url = self._get_best_image_url(img, base_url)
                        if url and url not in seen_urls:
                            seen_urls.add(url)
                            if text_index not in media_map:
                                media_map[text_index] = []
                            media_map[text_index].append({
                                "type": "img",
                                "data": url
                            })

                elif el.name == 'img':
                    # Пропускаємо img, якщо він всередині figure/picture
                    # (вже оброблений)
                    parent = el.parent
                    if parent and parent.name in ['figure', 'picture']:
                        continue
                    url = self._get_best_image_url(el, base_url)
                    if url and url not in seen_urls:
                        seen_urls.add(url)
                        if text_index not in media_map:
                            media_map[text_index] = []
                        media_map[text_index].append({
                            "type": "img",
                            "data": url
                        })

                elif el.name == 'video':
                    vid_url = el.get('src')
                    if not vid_url:
                        source = el.find('source')
                        if source:
                            vid_url = source.get('src')
                    vid_url = self.resolve_url(vid_url, base_url)
                    poster = self.resolve_url(
                        el.get('poster', ''), base_url
                    )
                    if vid_url and vid_url.startswith('http') and vid_url not in seen_urls:
                        seen_urls.add(vid_url)
                        if text_index not in media_map:
                            media_map[text_index] = []
                        media_map[text_index].append({
                            "type": "video",
                            "data": vid_url,
                            "poster": poster
                        })

                elif el.name == 'iframe':
                    src = el.get('src', '')
                    src = self.resolve_url(src, base_url)
                    if src and src.startswith('http') and src not in seen_urls:
                        # Перевіряємо, що це відео-iframe
                        video_hosts = [
                            'youtube.com', 'youtu.be', 'vimeo.com',
                            'dailymotion.com', 'rumble.com',
                        ]
                        if any(host in src.lower() for host in video_hosts):
                            seen_urls.add(src)
                            if text_index not in media_map:
                                media_map[text_index] = []
                            media_map[text_index].append({
                                "type": "video",
                                "data": src,
                                "poster": ""
                            })
            except Exception:
                continue

        return media_map

    def _build_content_list(self, raw_html, soup, base_url, cover_image=""):
        """
        Головна функція побудови content_list.
        Каскадно використовує trafilatura → readability → власну евристику.
        Завжди додає медіа з оригінального DOM.
        """
        download_images = self.state.get("download_images", True)
        content_list = []
        text_blocks = []
        used_readability = False

        # ═══ Спроба 1: trafilatura ═══
        traf_blocks = self._extract_with_trafilatura(raw_html, base_url)
        if traf_blocks and len(traf_blocks) >= 3:
            text_blocks = traf_blocks
            used_readability = True

        # ═══ Спроба 2: readability-lxml ═══
        if not text_blocks:
            read_soup = self._extract_with_readability(raw_html, base_url)
            if read_soup:
                text_blocks = self._soup_to_text_blocks(read_soup)
                if text_blocks and len(text_blocks) >= 3:
                    used_readability = True

        # ═══ Спроба 3: Власна евристика ═══
        if not text_blocks or len(text_blocks) < 3:
            cleaned_soup = self._deep_clean_html(deepcopy(soup))
            container = self._find_article_container(cleaned_soup)
            if container:
                text_blocks = self._soup_to_text_blocks(container)

        # Фільтруємо спам із текстових блоків
        text_blocks = [
            block for block in text_blocks
            if not self._is_spam_text(block.get("data", ""))
        ]

        if not text_blocks:
            return []

        # ═══ Додаємо медіа ═══
        if download_images:
            # Витягуємо медіа з оригінального (неочищеного) DOM
            # для кращого покриття
            cleaned_for_media = self._deep_clean_html(deepcopy(soup))
            media_container = self._find_article_container(cleaned_for_media)
            if not media_container:
                media_container = cleaned_for_media

            media_map = self._extract_media_from_container(
                media_container, base_url, cover_image
            )

            # Зшиваємо текст і медіа в хронологічному порядку
            for i, block in enumerate(text_blocks):
                # Вставляємо медіа ПЕРЕД поточним текстовим блоком
                if i in media_map:
                    for media_item in media_map[i]:
                        content_list.append(media_item)

                content_list.append(block)

            # Медіа після останнього текстового блоку
            last_idx = len(text_blocks)
            for idx in sorted(media_map.keys()):
                if idx >= last_idx:
                    for media_item in media_map[idx]:
                        content_list.append(media_item)
        else:
            content_list = text_blocks

        return content_list

    def _soup_to_text_blocks(self, container):
        """
        Конвертує BeautifulSoup контейнер у список текстових блоків.
        Пропускає медіа (вони обробляються окремо).
        """
        blocks = []
        seen_texts = set()

        for el in container.find_all(
                ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote',
                 'pre', 'td']
        ):
            try:
                text = el.get_text(strip=True)
                if not text or len(text) < MIN_TEXT_LENGTH:
                    continue

                # Дедуплікація
                text_hash = hashlib.md5(text.encode()).hexdigest()
                if text_hash in seen_texts:
                    continue
                seen_texts.add(text_hash)

                # Визначаємо тип
                tag_name = el.name
                if tag_name == 'h1':
                    tag_name = 'h2'  # h1 зазвичай заголовок сторінки
                elif tag_name in ['h5', 'h6']:
                    tag_name = 'h4'
                elif tag_name == 'pre':
                    tag_name = 'blockquote'
                elif tag_name == 'td':
                    tag_name = 'p'

                blocks.append({
                    "type": tag_name,
                    "data": text
                })
            except Exception:
                continue

        return blocks

    # ═══════════════════════════════════════════════════════════
    #                 ВІДКРИТТЯ ФАЙЛУ
    # ═══════════════════════════════════════════════════════════

    def open_saved_file(self, path):
        """Відкриває збережений файл у системному переглядачі"""
        try:
            if platform.system() == 'Darwin':
                subprocess.call(('open', path))
            elif platform.system() == 'Windows':
                os.startfile(path)
            else:
                subprocess.call(('xdg-open', path))
        except Exception:
            pass

    # ═══════════════════════════════════════════════════════════
    #              ГОЛОВНИЙ ЦИКЛ ОБРОБКИ (run)
    # ═══════════════════════════════════════════════════════════

    def run(self):
        """
        Головний метод QThread.
        Обробляє список URL: парсинг → переклад → генерація DOCX/PDF.
        """
        driver = None
        try:
            driver = self.get_browser_driver()
            total_urls = len(self.urls)
            is_single_url = (total_urls == 1)

            for index, url in enumerate(self.urls):
                if self._cancelled:
                    break

                # ── Пасхалка: Вікіпедія ──
                if "wikipedia.org" in url.lower():
                    self.status_updated.emit(
                        "Братику, Вікіпедія і так безкоштовна. "
                        "Що ти намагаєшся зробити? 🧠",
                        Theme.GOLD
                    )
                    time.sleep(2)
                    continue

                # ── Пасхалка: .ua на українську ──
                if ((".ua/" in url.lower() or ".com.ua/" in url.lower()) and
                        self.state.get("target_lang_name") == "Українська"):
                    current_status = (
                        "Перекладаю з української на... ще кращу "
                        "і солов'їнішу! 🇺🇦"
                    )
                else:
                    current_status = (
                        self.t("status_single_start") if is_single_url
                        else self.t("status_magic").format(index + 1, total_urls)
                    )

                self.status_updated.emit(current_status, Theme.GOLD)

                # ── Завантаження сторінки ──
                driver.get(url)
                try:
                    WebDriverWait(driver, 15).until(
                        EC.presence_of_element_located((By.TAG_NAME, "body"))
                    )
                except Exception:
                    pass
                time.sleep(5)

                if self._cancelled:
                    break

                raw_html = driver.page_source
                base_url = url

                # ── Парсинг ──
                try:
                    soup = BeautifulSoup(raw_html, "lxml")
                except Exception:
                    soup = BeautifulSoup(raw_html, "html.parser")

                # ── Метадані ──
                meta = self.extract_metadata(soup, base_url)
                if not meta['title']:
                    title_tag = soup.find('title')
                    meta['title'] = (
                        title_tag.get_text().strip() if title_tag
                        else f"Архівна стаття {index + 1}"
                    )

                # ── Побудова content_list (УНІВЕРСАЛЬНИЙ ПАРСИНГ) ──
                content_list = self._build_content_list(
                    raw_html, soup, base_url, meta.get('cover_image', '')
                )

                if not content_list:
                    # Спробуємо ще раз з body
                    body = soup.find('body')
                    if body:
                        content_list = self._soup_to_text_blocks(body)
                        content_list = [
                            b for b in content_list
                            if not self._is_spam_text(b.get("data", ""))
                        ]

                if not content_list:
                    continue

                if self._cancelled:
                    break

                # ═══════════════════════════════════════════════
                #         СТВОРЕННЯ ДОКУМЕНТА (DOCX)
                # ═══════════════════════════════════════════════

                doc = WordDocument()
                download_images = self.state.get("download_images", True)
                total_words = sum(
                    len(item.get("data", "").split())
                    for item in content_list
                    if item["type"] not in ["img", "video"]
                )

                # ── Заголовок ──
                title_translated = self.translate_text(meta['title'])
                t_p = doc.add_paragraph()
                t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                t_run = t_p.add_run(title_translated)
                t_run.bold = True
                t_run.font.size = Pt(24)
                t_run.font.name = self.state["font_family"]

                # ── Підзаголовок ──
                if meta['subtitle']:
                    sub_translated = self.translate_text(meta['subtitle'])
                    sub_p = doc.add_paragraph()
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sub_run = sub_p.add_run(sub_translated)
                    sub_run.italic = True
                    sub_run.font.size = Pt(14)
                    sub_run.font.name = self.state["font_family"]

                # ── Автор / Дата ──
                if meta['author'] or meta['date']:
                    auth_parts = []
                    if meta['author']:
                        auth_parts.append(
                            f"✍️ Автор: {self.translate_text(meta['author'])}"
                        )
                    if meta['date']:
                        auth_parts.append(f"📅 Дата: {meta['date']}")
                    auth_p = doc.add_paragraph()
                    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    auth_run = auth_p.add_run(" | ".join(auth_parts))
                    auth_run.bold = True
                    auth_run.font.size = Pt(11)

                # ── Обкладинка ──
                if download_images and meta['cover_image']:
                    try:
                        resp = requests.get(
                            meta['cover_image'],
                            headers=self.http_headers,
                            timeout=15
                        )
                        if resp.status_code == 200:
                            img_stream = self.convert_image_for_docx(resp.content)
                            if img_stream:
                                img_p = doc.add_paragraph()
                                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                img_run = img_p.add_run()
                                img_run.add_picture(img_stream, width=Inches(6.0))
                    except Exception:
                        pass

                doc.add_paragraph()

                # ── Час читання ──
                if self.state.get("add_read_time", True):
                    reading_time = max(1, total_words // 150)
                    if reading_time > 100:
                        rt_text = (
                            f"⏱ Час читання: ~{reading_time} хв. "
                            "(Впевнений, що це не 'Війна і мир'? "
                            "Завари чаю ☕📖)"
                        )
                    else:
                        rt_text = f"⏱ Час читання: ~{reading_time} хв."
                    rt_p = doc.add_paragraph()
                    rt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rt_run = rt_p.add_run(rt_text)
                    rt_run.italic = True
                    rt_run.font.size = Pt(12)

                # ── Зміст (TOC) ──
                if self.state.get("add_toc", True):
                    headings = [
                        item for item in content_list
                        if item["type"] in ['h2', 'h3', 'h4']
                    ]
                    if headings:
                        doc.add_paragraph()
                        toc_p = doc.add_paragraph()
                        toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        toc_run = toc_p.add_run(self.t("doc_toc_title"))
                        toc_run.bold = True
                        toc_run.font.name = self.state["font_family"]

                        for h in headings:
                            if self._cancelled:
                                return
                            translated_h = self.translate_text(h["data"])
                            h["translated_data"] = translated_h
                            level = int(h["type"][1])
                            indent = Inches((level - 2) * 0.3)
                            toc_item = doc.add_paragraph()
                            toc_item.paragraph_format.left_indent = indent
                            toc_r = toc_item.add_run(f"• {translated_h}")
                            toc_r.font.name = self.state["font_family"]
                            toc_r.font.size = Pt(self.state["font_size"] - 2)

                        doc.add_paragraph()

                # ── Основний контент ──
                total_elements = len(content_list)
                url_hash = hashlib.md5(url.encode()).hexdigest()
                cache_file = os.path.join(
                    self.cache_dir, f"{url_hash}.json"
                )
                cached_data = []

                for i, item in enumerate(content_list):
                    if self._cancelled:
                        return

                    progress_pct = int(((i + 1) / total_elements) * 100)
                    progress_txt = self.t("status_progress").format(
                        f"{i + 1}/{total_elements}"
                    )
                    self.status_updated.emit(progress_txt, Theme.GOLD)
                    self.progress_updated.emit(progress_pct)

                    try:
                        # ── Текстові елементи ──
                        if item["type"] in ['p', 'li', 'blockquote',
                                            'h2', 'h3', 'h4']:
                            translated = item.get("translated_data")
                            if not translated:
                                translated = self.translate_text(item["data"])

                            # Кеш
                            cached_data.append({
                                "type": item["type"],
                                "text": translated
                            })
                            try:
                                with open(cache_file, "w", encoding="utf-8") as f:
                                    json.dump(
                                        cached_data, f, ensure_ascii=False
                                    )
                            except Exception:
                                pass

                            # Двомовний режим
                            if (self.state.get("bilingual_mode", False) and
                                    item["type"] == 'p'):
                                orig_p = doc.add_paragraph()
                                orig_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                orig_run = orig_p.add_run(item["data"])
                                orig_run.italic = True
                                orig_run.font.size = Pt(
                                    self.state["font_size"] - 2
                                )
                                orig_run.font.name = self.state["font_family"]

                            # Додаємо в документ
                            if item["type"] in ['h2', 'h3', 'h4']:
                                h_level = int(item["type"][1])
                                h_p = doc.add_heading(translated, level=h_level)
                                if h_p.runs:
                                    h_p.runs[0].font.name = self.state["font_family"]
                            elif item["type"] == 'li':
                                p = doc.add_paragraph(
                                    translated, style='List Bullet'
                                )
                                if p.runs:
                                    p.runs[0].font.name = self.state["font_family"]
                                    p.runs[0].font.size = Pt(
                                        self.state["font_size"]
                                    )
                            elif item["type"] == 'blockquote':
                                try:
                                    p = doc.add_paragraph(
                                        translated, style='Intense Quote'
                                    )
                                except Exception:
                                    p = doc.add_paragraph()
                                    r = p.add_run(f"« {translated} »")
                                    r.italic = True
                                if p.runs:
                                    p.runs[0].font.name = self.state["font_family"]
                                    p.runs[0].font.size = Pt(
                                        self.state["font_size"]
                                    )
                            else:  # p
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                run = p.add_run(translated)
                                run.font.size = Pt(self.state["font_size"])
                                run.font.name = self.state["font_family"]

                        # ── Зображення ──
                        elif item["type"] == "img" and download_images:
                            try:
                                resp = requests.get(
                                    item["data"],
                                    headers=self.http_headers,
                                    timeout=15
                                )
                                if resp.status_code == 200:
                                    img_stream = self.convert_image_for_docx(
                                        resp.content
                                    )
                                    if img_stream:
                                        img_p = doc.add_paragraph()
                                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        img_run = img_p.add_run()
                                        img_run.add_picture(
                                            img_stream, width=Inches(5.5)
                                        )
                            except Exception:
                                pass

                        # ── Відео ──
                        elif item["type"] == "video" and download_images:
                            try:
                                # Постер відео
                                if item.get("poster"):
                                    try:
                                        resp = requests.get(
                                            item["poster"],
                                            headers=self.http_headers,
                                            timeout=10
                                        )
                                        if resp.status_code == 200:
                                            img_stream = self.convert_image_for_docx(
                                                resp.content
                                            )
                                            if img_stream:
                                                img_p = doc.add_paragraph()
                                                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                img_run = img_p.add_run()
                                                img_run.add_picture(
                                                    img_stream,
                                                    width=Inches(5.0)
                                                )
                                    except Exception:
                                        pass

                                # Посилання на відео
                                vid_p = doc.add_paragraph()
                                vid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                vid_run = vid_p.add_run(
                                    f"{self.t('video_link_text')}\n{item['data']}"
                                )
                                vid_run.font.size = Pt(11)
                                vid_run.underline = True
                            except Exception:
                                pass

                    except Exception:
                        continue

                # ── Метадані (футер документа) ──
                if self.state.get("add_metadata", True):
                    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    meta_p = doc.add_paragraph()
                    meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    meta_text = self.t("metadata_text").format(
                        url, current_date
                    )
                    meta_run = meta_p.add_run(meta_text)
                    meta_run.font.size = Pt(10)
                    meta_run.italic = True

                    # Пасхалка: пізня година
                    hour = datetime.now().hour
                    if 2 <= hour < 4:
                        meta_run.add_text(
                            "\n(P.S. Справжні генії працюють вночі, "
                            "але вам, мабуть, час іти спати... 🦉)"
                        )

                # ── Збереження файлу ──
                safe_title = re.sub(
                    r'[\\/:*?"<>|]', "", title_translated
                )[:80]
                base_path = os.path.join(
                    self.state["save_path"], safe_title
                )
                full_path = f"{base_path}.docx"

                counter = 1
                while True:
                    try:
                        doc.save(full_path)
                        break
                    except PermissionError:
                        full_path = f"{base_path} ({counter}).docx"
                        counter += 1

                # Видаляємо кеш
                if os.path.exists(cache_file):
                    try:
                        os.remove(cache_file)
                    except Exception:
                        pass

                # ── PDF конвертація ──
                final_path = full_path
                if self.state.get("output_format") == "pdf":
                    pdf_path = full_path.replace(".docx", ".pdf")
                    self.status_updated.emit(
                        self.t("status_pdf"), Theme.GOLD
                    )
                    try:
                        if docx2pdf_convert:
                            docx2pdf_convert(full_path, pdf_path)
                            if os.path.exists(full_path):
                                os.remove(full_path)
                            final_path = pdf_path
                    except Exception:
                        pass

                # ── Автовідкриття ──
                if self.state.get("auto_open", True):
                    self.open_saved_file(final_path)

                self.finished_success.emit(final_path)

            # ── Завершення ──
            if not self._cancelled:
                self.status_updated.emit(
                    self.t("status_success"), Theme.GREEN
                )
                self.progress_updated.emit(100)
                self.notification_requested.emit(
                    "Скарбниця Знань",
                    "✅ Обробку успішно завершено! Файл збережено."
                )

        except Exception as e:
            if not self._cancelled:
                self.error_occurred.emit(str(e))
                self.status_updated.emit(
                    self.t("status_error"), Theme.RED
                )

        finally:
            if driver:
                try:
                    driver.quit()
                except Exception:
                    pass
            self.finished_all.emit()