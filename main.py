"""
╔══════════════════════════════════════════════════════════════╗
║        СКАРБНИЦЯ ЗНАНЬ v6.0 — ПОВНА ВЕРСІЯ (PyQt6)         ║
║                                                              ║
║  Frontend (PyQt6) + Backend (Selenium, BeautifulSoup, DOCX) ║
║  Один файл. Готовий до запуску.                              ║
╚══════════════════════════════════════════════════════════════╝
"""

# ═══════════════════════════════════════════════════════════════
#                        ІМПОРТИ
# ═══════════════════════════════════════════════════════════════

import sys
import os
import re
import time
import json
import platform
import subprocess
import webbrowser
import requests
import hashlib
import warnings
import threading
from io import BytesIO
import urllib.parse
from datetime import datetime
from functools import partial

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QTextEdit, QProgressBar, QStackedWidget,
    QScrollArea, QFrame, QComboBox, QSlider, QFileDialog,
    QCheckBox, QMessageBox, QGraphicsDropShadowEffect,
    QSizePolicy, QSpacerItem, QSystemTrayIcon, QMenu
)
from PyQt6.QtCore import (
    Qt, QTimer, QPropertyAnimation, QEasingCurve,
    QPoint, QSize, pyqtSignal, QRect, QEvent, QThread
)
from PyQt6.QtGui import (
    QFont, QColor, QIcon, QPalette, QLinearGradient,
    QPainter, QBrush, QPen, QAction, QCursor, QPixmap
)

# ── Selenium ──
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# ── Парсинг та обробка ──
from bs4 import BeautifulSoup
from PIL import Image, ImageDraw

# ── Документи ──
from docx import Document as WordDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Переклад ──
from deep_translator import MicrosoftTranslator, GoogleTranslator

# ── PDF конвертація ──
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

# ── Буфер обміну ──
try:
    import pyperclip
except ImportError:
    pyperclip = None

warnings.filterwarnings("ignore", category=UserWarning, module='requests')

# ═══════════════════════════════════════════════════════════════
#              DPI AWARENESS (Windows 10/11)
# ═══════════════════════════════════════════════════════════════

try:
    if platform.system() == "Windows":
        import ctypes
        ctypes.windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    pass


# ═══════════════════════════════════════════════════════════════
#                    КОНСТАНТИ СТИЛІВ
# ═══════════════════════════════════════════════════════════════

class Theme:
    """Централізоване управління кольорами та шрифтами"""

    # ── Основна палітра ──
    GOLD = "#d4af37"
    GOLD_HOVER = "#b5952f"
    GOLD_DARK = "#8B7530"

    GREEN = "#2d5a27"
    GREEN_HOVER = "#1e3d1a"
    GREEN_LIGHT = "#3d7a37"

    RED = "#8b0000"
    RED_HOVER = "#4a0000"
    RED_LIGHT = "#c0392b"

    BLUE = "#2980b9"
    BLUE_HOVER = "#1c5980"

    ORANGE = "#d35400"
    ORANGE_HOVER = "#e67e22"

    # ── Фони (Dark Mode) ──
    BG_PRIMARY = "#0d0d0d"
    BG_SECONDARY = "#1a1a1a"
    BG_CARD = "#1e1e1e"
    BG_INPUT = "#2a2a2a"
    BG_HOVER = "#333333"

    # ── Текст ──
    TEXT_PRIMARY = "#f0f0f0"
    TEXT_SECONDARY = "#999999"
    TEXT_MUTED = "#666666"

    # ── Бордери ──
    BORDER = "#3a3a3a"
    BORDER_GOLD = "#d4af37"
    BORDER_FOCUS = "#e6c565"

    # ── Шрифти ──
    FONT_TITLE = "Georgia"
    FONT_UI = "Segoe UI"
    FONT_MONO = "Cascadia Code"


# ═══════════════════════════════════════════════════════════════
#                   ГЛОБАЛЬНІ СТИЛІ (QSS)
# ═══════════════════════════════════════════════════════════════

GLOBAL_STYLESHEET = f"""
    /* ── Загальне ── */
    QMainWindow {{
        background-color: {Theme.BG_PRIMARY};
    }}

    QWidget {{
        color: {Theme.TEXT_PRIMARY};
        font-family: "{Theme.FONT_UI}";
    }}

    /* ── Скролбари ── */
    QScrollBar:vertical {{
        background: {Theme.BG_SECONDARY};
        width: 8px;
        border-radius: 4px;
        margin: 0;
    }}
    QScrollBar::handle:vertical {{
        background: {Theme.GOLD_DARK};
        border-radius: 4px;
        min-height: 30px;
    }}
    QScrollBar::handle:vertical:hover {{
        background: {Theme.GOLD};
    }}
    QScrollBar::add-line:vertical,
    QScrollBar::sub-line:vertical {{
        height: 0px;
    }}
    QScrollBar::add-page:vertical,
    QScrollBar::sub-page:vertical {{
        background: none;
    }}

    /* ── Текстові поля ── */
    QTextEdit {{
        background-color: {Theme.BG_INPUT};
        color: {Theme.TEXT_PRIMARY};
        border: 2px solid {Theme.BORDER};
        border-radius: 16px;
        padding: 16px;
        font-size: 15px;
        font-family: "{Theme.FONT_UI}";
        selection-background-color: {Theme.GOLD_DARK};
    }}
    QTextEdit:focus {{
        border-color: {Theme.GOLD};
    }}

    /* ── Комбобокси ── */
    QComboBox {{
        background-color: {Theme.BG_INPUT};
        color: {Theme.TEXT_PRIMARY};
        border: 2px solid {Theme.BORDER};
        border-radius: 12px;
        padding: 10px 16px;
        font-size: 14px;
        min-width: 280px;
    }}
    QComboBox:hover {{
        border-color: {Theme.GOLD};
    }}
    QComboBox::drop-down {{
        border: none;
        width: 30px;
    }}
    QComboBox::down-arrow {{
        image: none;
        border: none;
    }}
    QComboBox QAbstractItemView {{
        background-color: {Theme.BG_CARD};
        color: {Theme.TEXT_PRIMARY};
        border: 1px solid {Theme.BORDER};
        border-radius: 8px;
        selection-background-color: {Theme.GOLD_DARK};
        padding: 4px;
    }}

    /* ── Слайдери ── */
    QSlider {{
        min-height: 30px;
    }}
    QSlider::groove:horizontal {{
        border: none;
        height: 6px;
        background: {Theme.BORDER};
        border-radius: 3px;
    }}
    QSlider::handle:horizontal {{
        background: {Theme.GOLD};
        border: 2px solid {Theme.GOLD_DARK};
        width: 20px;
        height: 20px;
        margin: -7px 0;
        border-radius: 10px;
    }}
    QSlider::handle:horizontal:hover {{
        background: {Theme.GOLD_HOVER};
        border-color: {Theme.GOLD};
    }}
    QSlider::sub-page:horizontal {{
        background: {Theme.GOLD};
        border-radius: 3px;
    }}

    /* ── Чекбокси ── */
    QCheckBox {{
        spacing: 12px;
        font-size: 14px;
        color: {Theme.TEXT_PRIMARY};
        padding: 6px 0;
    }}
    QCheckBox::indicator {{
        width: 22px;
        height: 22px;
        border-radius: 6px;
        border: 2px solid {Theme.BORDER};
        background-color: {Theme.BG_INPUT};
    }}
    QCheckBox::indicator:checked {{
        background-color: {Theme.GOLD};
        border-color: {Theme.GOLD};
    }}
    QCheckBox::indicator:hover {{
        border-color: {Theme.GOLD};
    }}

    /* ── Прогресбар ── */
    QProgressBar {{
        background-color: {Theme.BG_INPUT};
        border: none;
        border-radius: 6px;
        height: 12px;
        text-align: center;
        font-size: 0px;
    }}
    QProgressBar::chunk {{
        background: qlineargradient(
            x1:0, y1:0, x2:1, y2:0,
            stop:0 {Theme.GOLD_DARK},
            stop:1 {Theme.GOLD}
        );
        border-radius: 6px;
    }}

    /* ── Скрол-область ── */
    QScrollArea {{
        border: none;
        background-color: transparent;
    }}
    QScrollArea > QWidget > QWidget {{
        background-color: transparent;
    }}

    /* ── Тултіпи ── */
    QToolTip {{
        background-color: {Theme.BG_CARD};
        color: {Theme.GOLD};
        border: 1px solid {Theme.GOLD_DARK};
        border-radius: 8px;
        padding: 8px 12px;
        font-size: 13px;
    }}
"""


# ═══════════════════════════════════════════════════════════════
#               КАСТОМНІ КОМПОНЕНТИ (КНОПКИ)
# ═══════════════════════════════════════════════════════════════

class NoScrollComboBox(QComboBox):
    """Комбобокс, який не змінює значення при прокрутці колесом миші"""
    def wheelEvent(self, event):
        event.ignore()


class NoScrollSlider(QSlider):
    """Слайдер, який не змінює значення при прокрутці колесом миші"""
    def wheelEvent(self, event):
        event.ignore()


class GoldButton(QPushButton):
    """Кнопка з золотим бордером — для другорядних дій"""

    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.setFixedHeight(48)
        self.setFont(QFont(Theme.FONT_UI, 14, QFont.Weight.Bold))
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                color: {Theme.GOLD};
                border: 2px solid {Theme.GOLD};
                border-radius: 14px;
                padding: 8px 36px;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background-color: {Theme.GOLD};
                color: {Theme.BG_PRIMARY};
            }}
            QPushButton:pressed {{
                background-color: {Theme.GOLD_HOVER};
            }}
        """)


class PrimaryButton(QPushButton):
    """Головна зелена кнопка — для основних дій"""

    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.setFixedHeight(72)
        self.setFont(QFont(Theme.FONT_UI, 20, QFont.Weight.Bold))
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: {Theme.GREEN};
                color: white;
                border: none;
                border-radius: 20px;
                padding: 12px 40px;
                font-weight: bold;
                font-size: 20px;
            }}
            QPushButton:hover {{
                background-color: {Theme.GREEN_LIGHT};
            }}
            QPushButton:pressed {{
                background-color: {Theme.GREEN_HOVER};
            }}
            QPushButton:disabled {{
                background-color: {Theme.BG_HOVER};
                color: {Theme.TEXT_MUTED};
            }}
        """)

    def add_shadow(self):
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(25)
        shadow.setColor(QColor(45, 90, 39, 120))
        shadow.setOffset(0, 4)
        self.setGraphicsEffect(shadow)


class DangerButton(QPushButton):
    """Червона кнопка — для скасування"""

    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.setFixedHeight(72)
        self.setFont(QFont(Theme.FONT_UI, 16, QFont.Weight.Bold))
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                color: {Theme.RED};
                border: 2px solid {Theme.RED};
                border-radius: 20px;
                padding: 12px 32px;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background-color: {Theme.RED};
                color: white;
            }}
            QPushButton:pressed {{
                background-color: {Theme.RED_HOVER};
                color: white;
            }}
            QPushButton:disabled {{
                border-color: {Theme.TEXT_MUTED};
                color: {Theme.TEXT_MUTED};
            }}
        """)


class IconButton(QPushButton):
    """Іконка-кнопка без фону (для ⚙️, ← тощо)"""

    def __init__(self, text="", size=40, font_size=28, parent=None):
        super().__init__(text, parent)
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.setFixedSize(size, size)
        self.setFont(QFont("Arial", font_size))
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                color: {Theme.GOLD};
                border: none;
                border-radius: {size // 2}px;
            }}
            QPushButton:hover {{
                background-color: {Theme.BG_HOVER};
            }}
        """)


class FeatureButton(QPushButton):
    """Кнопка для меню "Про додаток" з кольоровим фоном"""

    def __init__(self, text="", bg_color="", hover_color="", parent=None):
        super().__init__(text, parent)
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.setFixedHeight(50)
        self.setFixedWidth(260)
        self.setFont(QFont(Theme.FONT_UI, 14, QFont.Weight.Bold))
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: {bg_color};
                color: white;
                border: none;
                border-radius: 16px;
                padding: 10px 24px;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background-color: {hover_color};
            }}
        """)


class CardFrame(QFrame):
    """Карточний контейнер з заокругленими кутами"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setStyleSheet(f"""
            QFrame {{
                background-color: {Theme.BG_CARD};
                border: 1px solid {Theme.BORDER};
                border-radius: 16px;
            }}
        """)


class SectionLabel(QLabel):
    """Заголовок секції в налаштуваннях"""

    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setFont(QFont(Theme.FONT_UI, 15, QFont.Weight.Bold))
        self.setStyleSheet(f"color: {Theme.TEXT_PRIMARY}; padding: 4px 0;")


class GoldSectionLabel(QLabel):
    """Золотий заголовок секції"""

    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setFont(QFont(Theme.FONT_UI, 15, QFont.Weight.Bold))
        self.setStyleSheet(f"color: {Theme.GOLD}; padding: 8px 0;")


class MutedLabel(QLabel):
    """Сірий підпис"""

    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setFont(QFont(Theme.FONT_UI, 12))
        self.setStyleSheet(f"color: {Theme.TEXT_MUTED};")


# ═══════════════════════════════════════════════════════════════
#            КАСТОМНИЙ TOGGLE SWITCH (замість QCheckBox)
# ═══════════════════════════════════════════════════════════════

class ToggleSwitch(QWidget):
    """Красивий перемикач (Toggle Switch) у стилі iOS/macOS"""
    toggled = pyqtSignal(bool)

    def __init__(self, text="", checked=False, parent=None):
        super().__init__(parent)
        self._checked = checked
        self._animation_pos = 26.0 if checked else 4.0

        self.setFixedHeight(40)
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))

        # Головний шар (горизонтальний)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)  # Відстань між повзунком і текстом

        # 1. Створюємо окремий віджет ТІЛЬКИ для малювання повзунка
        self.switch_visual = QWidget()
        self.switch_visual.setFixedSize(50, 28)
        # ВАЖЛИВО: Забороняємо глобальним стилям ламати наш малюнок
        self.switch_visual.setAttribute(Qt.WidgetAttribute.WA_StyledBackground, False)
        # Перевизначаємо малювання
        self.switch_visual.paintEvent = self._paint_switch

        # 2. Текст (іконка + опис)
        self.label = QLabel(text)
        self.label.setFont(QFont(Theme.FONT_UI, 14))
        self.label.setStyleSheet(f"color: {Theme.TEXT_PRIMARY};")

        # Додаємо усе на екран
        layout.addWidget(self.switch_visual)
        layout.addWidget(self.label)
        layout.addStretch()

    def _paint_switch(self, event):
        """Малює сам повзунок"""
        painter = QPainter(self.switch_visual)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)

        # Малюємо фон (Жовтий або Сірий)
        if self._checked:
            painter.setBrush(QBrush(QColor(Theme.GOLD)))
        else:
            painter.setBrush(QBrush(QColor(Theme.BORDER)))  # Сірий колір

        painter.setPen(Qt.PenStyle.NoPen)
        painter.drawRoundedRect(0, 2, 50, 24, 12, 12)

        # Малюємо білий кружок
        painter.setBrush(QBrush(QColor("white")))
        painter.drawEllipse(int(self._animation_pos), 5, 18, 18)

        painter.end()

    def mousePressEvent(self, event):
        """Клік по перемикачу"""
        self._checked = not self._checked
        target = 26.0 if self._checked else 4.0
        self._start_animation(target)
        self.toggled.emit(self._checked)
        super().mousePressEvent(event)

    def _start_animation(self, target):
        """Плавна анімація переміщення кружечка"""
        start = self._animation_pos
        step = (target - start) / 8.0
        self._timer = QTimer(self)

        def animate():
            self._animation_pos += step
            if (step > 0 and self._animation_pos >= target) or \
                    (step < 0 and self._animation_pos <= target):
                self._animation_pos = target
                self._timer.stop()
            self.switch_visual.update()  # Оновлюємо малюнок

        self._timer.timeout.connect(animate)
        self._timer.start(16)

    def isChecked(self):
        return self._checked

    def setChecked(self, val):
        self._checked = val
        self._animation_pos = 26.0 if val else 4.0
        self.switch_visual.update()


# ═══════════════════════════════════════════════════════════════
#          SCRAPING WORKER (QThread замість threading)
# ═══════════════════════════════════════════════════════════════

class ScrapingWorker(QThread):
    """
    Воркер для парсингу статей у фоновому потоці.
    Спілкується з головним вікном через сигнали.
    """
    progress_updated = pyqtSignal(int)           # Прогрес 0-100
    status_updated = pyqtSignal(str, str)        # (текст, колір)
    finished_success = pyqtSignal(str)           # Шлях до файлу
    finished_all = pyqtSignal()                  # Всі URL оброблено
    error_occurred = pyqtSignal(str)             # Повідомлення про помилку
    notification_requested = pyqtSignal(str, str)  # (title, message)

    def __init__(self, urls, state, locales, parent=None):
        super().__init__(parent)
        self.urls = urls
        self.state = state.copy()
        self.locales = locales
        self._cancelled = False

        self.http_headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                          '(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8'
        }

        self.translation_languages = {
            "Українська": "uk", "English": "en", "Polski": "pl",
            "Deutsch": "de", "Français": "fr", "Español": "es"
        }

        self.cache_dir = "backup_cache"
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)

    def cancel(self):
        self._cancelled = True

    def t(self, key):
        lang = self.state.get("ui_language", "uk")
        return self.locales.get(lang, self.locales["uk"]).get(key, key)

    # ── Браузер ──
    def get_browser_driver(self):
        prefs = {"profile.managed_default_content_settings.javascript": 2}
        try:
            options = EdgeOptions()
            options.add_argument("--headless=new")
            options.add_argument("--disable-gpu")
            options.add_argument("--no-sandbox")
            options.add_argument("--disable-dev-shm-usage")
            options.add_experimental_option("prefs", prefs)
            return webdriver.Edge(options=options)
        except Exception:
            try:
                options = ChromeOptions()
                options.add_argument("--headless=new")
                options.add_argument("--disable-gpu")
                options.add_argument("--no-sandbox")
                options.add_argument("--disable-dev-shm-usage")
                options.add_experimental_option("prefs", prefs)
                return webdriver.Chrome(options=options)
            except Exception:
                raise Exception("Не вдалося запустити ні Edge, ні Chrome браузер на цьому ПК.")

    # ── Переклад ──
    def translate_text(self, text):
        if not text or len(text.strip()) < 5:
            return text
        target_name = self.state.get("target_lang_name", "Українська")
        target_code = self.translation_languages.get(target_name, "uk")
        engine = self.state.get("translation_engine", "Google Translator")

        if engine == "Microsoft Translator":
            try:
                return MicrosoftTranslator(target=target_code).translate(text)
            except Exception:
                try:
                    return GoogleTranslator(source='auto', target=target_code).translate(text)
                except Exception:
                    return text
        else:
            try:
                return GoogleTranslator(source='auto', target=target_code).translate(text)
            except Exception:
                try:
                    return MicrosoftTranslator(target=target_code).translate(text)
                except Exception:
                    return text

    # ── Утиліти зображень ──
    def convert_image_for_docx(self, image_bytes):
        try:
            img = Image.open(BytesIO(image_bytes))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            elif img.mode != "RGB":
                img = img.convert("RGB")
            output = BytesIO()
            img.save(output, format="JPEG")
            output.seek(0)
            return output
        except Exception:
            return None

    def resolve_url(self, url, base_url):
        if not url:
            return ""
        if url.startswith('data:image'):
            return ""
        if url.startswith('http'):
            return url
        if url.startswith('//'):
            return 'https:' + url
        return urllib.parse.urljoin(base_url, url)

    # ── Метадані ──
    def extract_metadata(self, soup, base_url):
        meta = {"title": "", "subtitle": "", "author": "", "date": "", "cover_image": ""}

        def safe_get(d, key, default=""):
            if isinstance(d, dict):
                return d.get(key, default)
            return default

        try:
            og_img = soup.find('meta', property='og:image') or soup.find('meta', attrs={'name': 'twitter:image'})
            if og_img and og_img.get('content'):
                meta['cover_image'] = self.resolve_url(og_img['content'], base_url)
        except Exception:
            pass

        try:
            json_scripts = soup.find_all('script', type='application/ld+json')
            for script in json_scripts:
                if not script.string:
                    continue
                try:
                    raw_data = json.loads(script.string)
                    article_data = None
                    if isinstance(raw_data, list):
                        for item_d in raw_data:
                            if isinstance(item_d, dict) and safe_get(item_d, '@type') in [
                                'NewsArticle', 'Article', 'ReportageNewsArticle', 'WebPage'
                            ]:
                                article_data = item_d
                                break
                        if not article_data and len(raw_data) > 0 and isinstance(raw_data[0], dict):
                            article_data = raw_data[0]
                    elif isinstance(raw_data, dict):
                        article_data = raw_data
                    if not isinstance(article_data, dict):
                        continue

                    if '@graph' in article_data and isinstance(article_data['@graph'], list):
                        for graph_item in article_data['@graph']:
                            if isinstance(graph_item, dict) and safe_get(graph_item, '@type') in [
                                'NewsArticle', 'Article', 'WebPage'
                            ]:
                                article_data = graph_item
                                break

                    if not meta['title']:
                        meta['title'] = safe_get(article_data, 'headline') or safe_get(article_data, 'name')
                    if not meta['subtitle']:
                        meta['subtitle'] = safe_get(article_data, 'description')
                    if not meta['date']:
                        d_val = safe_get(article_data, 'datePublished')
                        meta['date'] = str(d_val).split('T')[0] if d_val else ''

                    author_data = safe_get(article_data, 'author')
                    if isinstance(author_data, list):
                        authors = [safe_get(a, 'name') for a in author_data if isinstance(a, dict)]
                        meta['author'] = ", ".join(filter(None, authors))
                    else:
                        meta['author'] = safe_get(author_data, 'name') if isinstance(author_data, dict) else str(
                            author_data or "")

                    if not meta['cover_image']:
                        img_data = safe_get(article_data, 'image')
                        raw_img = ""
                        if isinstance(img_data, list) and len(img_data) > 0:
                            raw_img = safe_get(img_data[0], 'url') if isinstance(img_data[0], dict) else str(
                                img_data[0])
                        elif isinstance(img_data, dict):
                            raw_img = safe_get(img_data, 'url')
                        elif isinstance(img_data, str):
                            raw_img = img_data
                        if raw_img:
                            meta['cover_image'] = self.resolve_url(raw_img, base_url)
                except Exception:
                    continue
        except Exception:
            pass

        try:
            if not meta['title']:
                t = soup.find('meta', property='og:title') or soup.find('meta', attrs={'name': 'title'})
                meta['title'] = t['content'] if (t and t.has_attr('content')) else ""
            if not meta['subtitle']:
                d = soup.find('meta', property='og:description') or soup.find('meta', attrs={'name': 'description'})
                meta['subtitle'] = d['content'] if (d and d.has_attr('content')) else ""
            if not meta['author']:
                a = soup.find('meta', attrs={'name': 'author'}) or soup.find('meta', property='article:author')
                meta['author'] = a['content'] if (a and a.has_attr('content')) else ""
        except Exception:
            pass

        return meta

    # ── Очищення HTML ──
    def clean_junk_html(self, soup):
        try:
            for unwanted in soup(['nav', 'aside', 'footer', 'header', 'script', 'style', 'button', 'form', 'svg']):
                try:
                    unwanted.decompose()
                except Exception:
                    pass

            junk_keywords = [
                'newsletter', 'promo', 'recirc', 'related', 'recommend', 'social',
                'share', 'author-bio', 'bottom', 'ad-', 'advertisement', 'trending'
            ]

            for tag in soup.find_all(['div', 'section', 'ul']):
                try:
                    c_list = tag.get('class', [])
                    class_str = ' '.join(c_list).lower() if isinstance(c_list, list) else str(c_list).lower()
                    id_str = str(tag.get('id', '')).lower()
                    if any(k in class_str or k in id_str for k in junk_keywords):
                        tag.decompose()
                except Exception:
                    continue
        except Exception:
            pass
        return soup

    # ── Відкриття файлу ──
    def open_saved_file(self, path):
        try:
            if platform.system() == 'Darwin':
                subprocess.call(('open', path))
            elif platform.system() == 'Windows':
                os.startfile(path)
            else:
                subprocess.call(('xdg-open', path))
        except Exception:
            pass

    # ── ГОЛОВНИЙ ЦИКЛ ОБРОБКИ ──
    def run(self):
        driver = None
        try:
            driver = self.get_browser_driver()
            total_urls = len(self.urls)
            is_single_url = (total_urls == 1)

            for index, url in enumerate(self.urls):
                if self._cancelled:
                    break

                # ── Пасхалка: Вікіпедія ──
                if "wikipedia.org" in url.lower():
                    self.status_updated.emit(
                        "Братику, Вікіпедія і так безкоштовна. Що ти намагаєшся зробити? 🧠",
                        Theme.GOLD
                    )
                    time.sleep(2)
                    continue

                # ── Пасхалка: .ua на українську ──
                if (".ua/" in url.lower() or ".com.ua/" in url.lower()) and \
                        self.state.get("target_lang_name") == "Українська":
                    current_status = "Перекладаю з української на... ще кращу і солов'їнішу! 🇺🇦"
                else:
                    current_status = self.t("status_single_start") if is_single_url else \
                        self.t("status_magic").format(index + 1, total_urls)

                self.status_updated.emit(current_status, Theme.GOLD)

                driver.get(url)
                try:
                    WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.TAG_NAME, "body"))
                    )
                except Exception:
                    pass
                time.sleep(5)

                if self._cancelled:
                    break

                raw_html = driver.page_source
                base_url = url

                try:
                    soup = BeautifulSoup(raw_html, "lxml")
                except Exception:
                    soup = BeautifulSoup(raw_html, "html.parser")

                meta = self.extract_metadata(soup, base_url)
                if not meta['title']:
                    title_tag = soup.find('title')
                    meta['title'] = title_tag.get_text().strip() if title_tag else f"Архівна стаття {index + 1}"

                article_container = soup.find('article')
                if not article_container:
                    article_container = soup.find('main')
                if not article_container:
                    article_container = soup.find('div', class_=re.compile(
                        r'article|story|post|content|main', re.I))
                if not article_container:
                    article_container = soup

                article_container = self.clean_junk_html(article_container)
                elements = article_container.find_all(
                    ['p', 'img', 'picture', 'video', 'iframe', 'h2', 'h3', 'h4', 'li', 'blockquote']
                )

                content_list = []
                total_words = 0
                download_images = self.state.get("download_images", True)
                spam_phrases = [
                    "підпишіться на", "subscribe to", "sign up for",
                    "read more:", "more from", "newsletter"
                ]

                for el in elements:
                    try:
                        if el.name in ['p', 'li', 'blockquote', 'h2', 'h3', 'h4']:
                            txt = el.get_text().strip()
                            if len(txt) < 150 and any(s in txt.lower() for s in spam_phrases):
                                continue
                            if len(txt) > 20 and "cookie" not in txt.lower() and "javascript" not in txt.lower():
                                content_list.append({"type": el.name, "data": txt})
                                total_words += len(txt.split())

                        elif el.name in ['img', 'picture'] and download_images:
                            if el.name == 'picture':
                                img_tag = el.find('img')
                                if img_tag:
                                    el = img_tag
                                else:
                                    continue

                            img_url = el.get('src') or el.get('data-src')
                            if not img_url and el.get('srcset'):
                                img_url = el.get('srcset').split(',')[0].split(' ')[0]

                            img_url = self.resolve_url(img_url, base_url)
                            if img_url and img_url.startswith('http') and img_url != meta['cover_image']:
                                content_list.append({"type": "img", "data": img_url})

                        elif el.name in ['video', 'iframe'] and download_images:
                            vid_url = el.get('src')
                            if el.name == 'video' and not vid_url:
                                source = el.find('source')
                                if source:
                                    vid_url = source.get('src')

                            vid_url = self.resolve_url(vid_url, base_url)
                            poster_url = self.resolve_url(el.get('poster'), base_url) if el.name == 'video' else ""

                            if vid_url and vid_url.startswith('http'):
                                content_list.append({"type": "video", "data": vid_url, "poster": poster_url})
                    except Exception:
                        continue

                if not content_list:
                    continue

                if self._cancelled:
                    break

                # ── Створення документа ──
                doc = WordDocument()

                title_translated = self.translate_text(meta['title'])
                t_p = doc.add_paragraph()
                t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                t_run = t_p.add_run(title_translated)
                t_run.bold = True
                t_run.font.size = Pt(24)
                t_run.font.name = self.state["font_family"]

                if meta['subtitle']:
                    sub_translated = self.translate_text(meta['subtitle'])
                    sub_p = doc.add_paragraph()
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sub_run = sub_p.add_run(sub_translated)
                    sub_run.italic = True
                    sub_run.font.size = Pt(14)
                    sub_run.font.name = self.state["font_family"]

                if meta['author'] or meta['date']:
                    auth_text = []
                    if meta['author']:
                        auth_text.append(f"✍️ Автор: {self.translate_text(meta['author'])}")
                    if meta['date']:
                        auth_text.append(f"📅 Дата: {meta['date']}")
                    auth_p = doc.add_paragraph()
                    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    auth_run = auth_p.add_run(" | ".join(auth_text))
                    auth_run.bold = True
                    auth_run.font.size = Pt(11)

                if download_images and meta['cover_image']:
                    try:
                        response = requests.get(meta['cover_image'], headers=self.http_headers, timeout=10)
                        if response.status_code == 200:
                            img_stream = self.convert_image_for_docx(response.content)
                            if img_stream:
                                img_p = doc.add_paragraph()
                                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                img_run = img_p.add_run()
                                img_run.add_picture(img_stream, width=Inches(6.0))
                    except Exception:
                        pass

                doc.add_paragraph()

                if self.state.get("add_read_time", True):
                    reading_time = max(1, total_words // 150)
                    if reading_time > 100:
                        rt_text = f"⏱ Час читання: ~{reading_time} хв. (Впевнений, що це не 'Війна і мир'? Завари чаю ☕📖)"
                    else:
                        rt_text = f"⏱ Час читання: ~{reading_time} хв."
                    rt_p = doc.add_paragraph()
                    rt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rt_run = rt_p.add_run(rt_text)
                    rt_run.italic = True
                    rt_run.font.size = Pt(12)

                if self.state.get("add_toc", True):
                    headings = [item_d for item_d in content_list if item_d["type"] in ['h2', 'h3', 'h4']]
                    if headings:
                        doc.add_paragraph()
                        toc_title_p = doc.add_paragraph()
                        toc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        toc_title_run = toc_title_p.add_run(self.t("doc_toc_title"))
                        toc_title_run.bold = True
                        toc_title_run.font.name = self.state["font_family"]

                        for h in headings:
                            if self._cancelled:
                                return
                            translated_h = self.translate_text(h["data"])
                            h["translated_data"] = translated_h
                            level = int(h["type"][1])
                            indent = Inches((level - 2) * 0.3)
                            toc_p = doc.add_paragraph()
                            toc_p.paragraph_format.left_indent = indent
                            toc_run = toc_p.add_run(f"• {translated_h}")
                            toc_run.font.name = self.state["font_family"]
                            toc_run.font.size = Pt(self.state["font_size"] - 2)
                        doc.add_paragraph()

                total_elements = len(content_list)
                url_hash = hashlib.md5(url.encode()).hexdigest()
                cache_file = os.path.join(self.cache_dir, f"{url_hash}.json")
                cached_data = []

                for i, item_d in enumerate(content_list):
                    if self._cancelled:
                        return

                    progress_pct = int(((i + 1) / total_elements) * 100)
                    progress_txt = self.t("status_progress").format(f"{i + 1}/{total_elements}")
                    self.status_updated.emit(progress_txt, Theme.GOLD)
                    self.progress_updated.emit(progress_pct)

                    try:
                        if item_d["type"] in ['p', 'li', 'blockquote', 'h2', 'h3', 'h4']:
                            translated = item_d.get("translated_data")
                            if not translated:
                                translated = self.translate_text(item_d["data"])
                            cached_data.append({"type": item_d["type"], "text": translated})
                            with open(cache_file, "w", encoding="utf-8") as f:
                                json.dump(cached_data, f, ensure_ascii=False)

                            if self.state.get("bilingual_mode", False) and item_d["type"] == 'p':
                                orig_p = doc.add_paragraph()
                                orig_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                orig_run = orig_p.add_run(item_d["data"])
                                orig_run.italic = True
                                orig_run.font.size = Pt(self.state["font_size"] - 2)
                                orig_run.font.name = self.state["font_family"]

                            if item_d["type"] in ['h2', 'h3', 'h4']:
                                h_level = int(item_d["type"][1])
                                h_p = doc.add_heading(translated, level=h_level)
                                h_p.runs[0].font.name = self.state["font_family"]
                            elif item_d["type"] == 'li':
                                p = doc.add_paragraph(translated, style='List Bullet')
                                p.runs[0].font.name = self.state["font_family"]
                                p.runs[0].font.size = Pt(self.state["font_size"])
                            elif item_d["type"] == 'blockquote':
                                p = doc.add_paragraph(translated, style='Intense Quote')
                                p.runs[0].font.name = self.state["font_family"]
                                p.runs[0].font.size = Pt(self.state["font_size"])
                            else:
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                run = p.add_run(translated)
                                run.font.size = Pt(self.state["font_size"])
                                run.font.name = self.state["font_family"]

                        elif item_d["type"] == "img":
                            try:
                                response = requests.get(item_d["data"], headers=self.http_headers, timeout=10)
                                if response.status_code == 200:
                                    img_stream = self.convert_image_for_docx(response.content)
                                    if img_stream:
                                        img_p = doc.add_paragraph()
                                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        img_run = img_p.add_run()
                                        img_run.add_picture(img_stream, width=Inches(5.5))
                            except Exception:
                                pass

                        elif item_d["type"] == "video":
                            try:
                                if item_d.get("poster"):
                                    response = requests.get(
                                        item_d["poster"], headers=self.http_headers, timeout=10
                                    )
                                    if response.status_code == 200:
                                        img_stream = self.convert_image_for_docx(response.content)
                                        if img_stream:
                                            img_p = doc.add_paragraph()
                                            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            img_run = img_p.add_run()
                                            img_run.add_picture(img_stream, width=Inches(5.0))

                                vid_p = doc.add_paragraph()
                                vid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                vid_run = vid_p.add_run(
                                    f"{self.t('video_link_text')}\n{item_d['data']}"
                                )
                                vid_run.font.size = Pt(11)
                                vid_run.underline = True
                            except Exception:
                                pass

                    except Exception:
                        continue

                if self.state.get("add_metadata", True):
                    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    meta_p = doc.add_paragraph()
                    meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    meta_text = self.t("metadata_text").format(url, current_date)
                    meta_run = meta_p.add_run(meta_text)
                    meta_run.font.size = Pt(10)
                    meta_run.italic = True

                    hour = datetime.now().hour
                    if 2 <= hour < 4:
                        meta_run.add_text(
                            "\n(P.S. Справжні генії працюють вночі, але вам, мабуть, час іти спати... 🦉)"
                        )

                safe_title = re.sub(r'[\\/:*?"<>|]', "", title_translated)[:80]
                base_path = os.path.join(self.state["save_path"], safe_title)
                full_path = f"{base_path}.docx"

                counter = 1
                while True:
                    try:
                        doc.save(full_path)
                        break
                    except PermissionError:
                        full_path = f"{base_path} ({counter}).docx"
                        counter += 1

                if os.path.exists(cache_file):
                    os.remove(cache_file)

                final_path = full_path
                if self.state.get("output_format") == "pdf":
                    pdf_path = full_path.replace(".docx", ".pdf")
                    self.status_updated.emit(self.t("status_pdf"), Theme.GOLD)
                    try:
                        if docx2pdf_convert:
                            docx2pdf_convert(full_path, pdf_path)
                            if os.path.exists(full_path):
                                os.remove(full_path)
                            final_path = pdf_path
                    except Exception:
                        pass

                if self.state.get("auto_open", True):
                    self.open_saved_file(final_path)

                self.finished_success.emit(final_path)

            if not self._cancelled:
                self.status_updated.emit(self.t("status_success"), Theme.GREEN)
                self.progress_updated.emit(100)
                self.notification_requested.emit(
                    "Скарбниця Знань",
                    "✅ Обробку успішно завершено! Файл збережено."
                )

        except Exception as e:
            if not self._cancelled:
                self.error_occurred.emit(str(e))
                self.status_updated.emit(self.t("status_error"), Theme.RED)

        finally:
            if driver:
                driver.quit()
            self.finished_all.emit()


# ═══════════════════════════════════════════════════════════════
#                     ГОЛОВНЕ ВІКНО
# ═══════════════════════════════════════════════════════════════

class TreasuryApp(QMainWindow):
    """Головне вікно додатку Скарбниця Знань"""

    def __init__(self):
        super().__init__()

        # ── Конфігурація вікна ──
        self.setWindowTitle("Скарбниця Знань v6.0")
        self.setMinimumSize(1000, 900)
        self.resize(1000, 900)

        try:
            self.setWindowIcon(QIcon("icon.ico"))
        except Exception:
            pass

        # ── Лічильники пасхалок ──
        self.temple_clicks = 0
        self.title_clicks = 0
        self.empty_clicks = 0
        self.about_clicks = 0
        self.theme_clicks = 0
        self.theme_click_time = 0
        self.format_clicks = 0
        self.format_click_time = 0
        self.bilingual_clicks = 0
        self.bilingual_click_time = 0
        self.font_slider_clicks = 0
        self.cancel_folder_clicks = 0
        self.hacker_mode = False
        self.start_process_time = 0

        # ── Воркер ──
        self.worker = None
        self.is_processing = False

        # ── Файл налаштувань ──
        self.config_file = "settings.json"

        # ── Стан додатку ──
        self.state = self._load_settings()

        # ── Локалізація (ПОВНА — uk + en) ──
        self.locales = {
            "uk": {
                "title": "Скарбниця Знань v6.0",
                "placeholder": "Вставте посилання (можна кілька, кожне з нового рядка)...",
                "status_wait": "Очікування посилань...",
                "btn_digitize": "ОЦИФРУВАТИ В АРХІВ",
                "btn_cancel": "СКАСУВАТИ",

                "settings_title": "Налаштування",
                "path_lbl": "📍 Папка збереження:",
                "btn_choose": "Обрати",
                "font_lbl": "🖋️ Шрифт:",
                "size_lbl": "📏 Розмір тексту:",
                "format_lbl": "📄 Формат збереження:",
                "engine_lbl": "🤖 Рушій перекладу:",
                "theme_lbl": "🌓 Тема інтерфейсу:",
                "theme_dark": "Темний режим",
                "ui_lang_lbl": "🌍 Мова інтерфейсу:",
                "target_lang_lbl": "🎯 Перекладати статтю на:",

                "additional_features_lbl": "🛠 Додаткові функції",
                "setting_read_time": "⏱ Додавати орієнтовний час читання",
                "setting_bilingual": "📖 Двомовний режим (Оригінал + переклад)",
                "setting_auto_open": "🚀 Автоматично відкривати документ",
                "setting_images": "🖼️ Завантажувати зображення та відео",
                "setting_toc": "📑 Додавати автоматичний Зміст",
                "setting_metadata": "🔗 Додавати посилання на джерело",
                "setting_tray_close": "⏬ Згортати в трей при закритті",

                "btn_about": "ℹ️ Про додаток",
                "about_title": "Про додаток",
                "about_desc": (
                    "Цей додаток створений для вільного доступу до інформації "
                    "без кордонів та обмежень.\n\n"
                    "Основна мета «Скарбниці Знань» — дати вам змогу читати "
                    "статті, журнали та новини, які заблоковані у вашій країні, "
                    "приховані за пейволом або вимагають платної підписки."
                ),
                "btn_features": "Основні функції ⭐️",
                "btn_how_it_works": "Як це працює ⚙️",
                "btn_donate": "☕ Підтримати автора",
                "btn_feedback": "@ Зворотний зв'язок",

                "how_it_works_title": "Механізм роботи",
                "how_it_works_text": (
                    "🔍 Парсинг:\nПрограма використовує 'Снайперський режим' для пошуку "
                    "тексту та витягує метадані (обкладинку, автора).\n\n"
                    "🧠 Аналіз:\nАлгоритм відкидає меню, рекламні банери та списки "
                    "'схожих новин'.\n\n"
                    "🌍 Переклад та Збірка:\nАбзаци пропускаються через перекладач і "
                    "зшиваються у документ разом із картинками."
                ),

                "premium_title": "Доступ до Premium",
                "premium_text": (
                    "Вітаю, шукачу ексклюзиву! 🎩\n\n"
                    "Ніякого 'Premium' у цьому додатку немає і ніколи не буде.\n\n"
                    "Користуйся на здоров'я, розширюй кругозір і нехай ця програма "
                    "служить тобі вірою і правдою."
                ),

                "features_title": "Можливості додатку",
                "features_text": (
                    "Повний перелік функцій Скарбниці Знань:\n\n"
                    "🔹 Згортання у Трей (робота у фоні)\n"
                    "🔹 Швидке викачування статті з буфера обміну\n"
                    "🔹 Обхід пейволів (читання платних статей)\n"
                    "🔹 Витягування обкладинки, відео, автора та дати\n"
                    "🔹 Збереження файлів у DOCX та PDF\n"
                    "🔹 Двомовний режим\n"
                    "🔹 Автоматичний Зміст для довгих статей\n"
                    "🔹 15+ пасхалок 🥚"
                ),
                "btn_back": "Повернутися",

                "status_single_start": "🌐 Завантаження та обробка статті...",
                "status_magic": "🌐 Старт пакетної обробки (Стаття {} з {})...",
                "status_progress": "📜 Обробка {} з елементів...",
                "status_success": "✅ Усі документи успішно збережено!",
                "status_error": "❌ Помилка обробки",
                "status_cancelled": "🛑 Процес скасовано",
                "status_pdf": "📄 Конвертація у PDF...",
                "msg_error_txt": "Текст або контент не знайдено.",
                "msg_invalid_url": "Знайдено некоректне посилання. Перевірте ввід.",
                "doc_toc_title": "--- ЗМІСТ ---",
                "metadata_text": "\n\n---\n🔗 Джерело: {}\n📅 Збережено: {}",
                "video_link_text": "▶️ [Дивитися прикріплене відео]",
                "tray_clipboard_err": "У буфері обміну немає посилання!",
                "tray_processing_err": "Програма вже зайнята обробкою!",

                "qm_btn_paste": "📥 Оцифрувати з буфера",
                "qm_btn_open": "🖥️ Відкрити програму",
                "qm_btn_exit": "❌ Вийти повністю",
                "qm_tray_open": "Відкрити меню"
            },
            "en": {
                "title": "Treasury of Knowledge v6.0",
                "placeholder": "Paste URLs here (one per line)...",
                "status_wait": "Waiting for URLs...",
                "btn_digitize": "DIGITIZE TO ARCHIVE",
                "btn_cancel": "CANCEL",

                "settings_title": "Settings",
                "path_lbl": "📍 Save Directory:",
                "btn_choose": "Browse",
                "font_lbl": "🖋️ Font:",
                "size_lbl": "📏 Text Size:",
                "format_lbl": "📄 Format:",
                "engine_lbl": "🤖 Translation Engine:",
                "theme_lbl": "🌓 Theme:",
                "theme_dark": "Dark Mode",
                "ui_lang_lbl": "🌍 UI Language:",
                "target_lang_lbl": "🎯 Target Language:",

                "additional_features_lbl": "🛠 Features",
                "setting_read_time": "⏱ Add estimated reading time",
                "setting_bilingual": "📖 Bilingual Mode (Original + Translation)",
                "setting_auto_open": "🚀 Auto-open document after creation",
                "setting_images": "🖼️ Download images, covers & videos",
                "setting_toc": "📑 Add Table of Contents (for long articles)",
                "setting_metadata": "🔗 Add source URL at the end of file",
                "setting_tray_close": "⏬ Minimize to tray on close (background mode)",

                "btn_about": "ℹ️ About",
                "about_title": "About Application",
                "about_desc": (
                    "This application was created for free access to information "
                    "without borders and restrictions.\n\n"
                    "The main goal of «Treasury of Knowledge» is to let you read "
                    "articles, magazines and news that are blocked in your country, "
                    "hidden behind a paywall, or require a paid subscription."
                ),
                "btn_features": "Features ⭐️",
                "btn_how_it_works": "How it works ⚙️",
                "btn_donate": "☕ Support Author",
                "btn_feedback": "@ Send Feedback",

                "how_it_works_title": "How It Works",
                "how_it_works_text": (
                    "🔍 Parsing:\nThe app uses 'Sniper Mode' to locate the article text "
                    "and extracts metadata (cover image, author).\n\n"
                    "🧠 Analysis:\nThe algorithm discards menus, ad banners, and "
                    "'related articles' lists.\n\n"
                    "🌍 Translation & Assembly:\nParagraphs are passed through the translator "
                    "and assembled into a document along with images."
                ),

                "premium_title": "Premium Access",
                "premium_text": (
                    "Greetings, seeker of exclusives! 🎩\n\n"
                    "There is no 'Premium' in this app and there never will be.\n\n"
                    "Use it freely, broaden your horizons and may this program "
                    "serve you faithfully!"
                ),

                "features_title": "App Features",
                "features_text": (
                    "Full list of Treasury of Knowledge features:\n\n"
                    "🔹 System Tray mode (background operation)\n"
                    "🔹 Quick article download from clipboard\n"
                    "🔹 Bypass paywalls (read paid articles)\n"
                    "🔹 Extract cover images, videos, author & date\n"
                    "🔹 Save files as DOCX and PDF\n"
                    "🔹 Bilingual mode\n"
                    "🔹 Automatic Table of Contents for long articles\n"
                    "🔹 15+ Easter Eggs 🥚"
                ),
                "btn_back": "Go Back",

                "status_single_start": "🌐 Loading and processing article...",
                "status_magic": "🌐 Batch processing (Article {} of {})...",
                "status_progress": "📜 Processing element {} ...",
                "status_success": "✅ All documents saved successfully!",
                "status_error": "❌ Processing error",
                "status_cancelled": "🛑 Cancelled",
                "status_pdf": "📄 Converting to PDF...",
                "msg_error_txt": "No text or content found.",
                "msg_invalid_url": "Invalid URL detected. Please check your input.",
                "doc_toc_title": "--- TABLE OF CONTENTS ---",
                "metadata_text": "\n\n---\n🔗 Source: {}\n📅 Saved: {}",
                "video_link_text": "▶️ [Watch attached video]",
                "tray_clipboard_err": "No URL found in clipboard!",
                "tray_processing_err": "App is already processing!",

                "qm_btn_paste": "📥 Digitize from clipboard",
                "qm_btn_open": "🖥️ Open Main App",
                "qm_btn_exit": "❌ Quit App",
                "qm_tray_open": "Open Menu"
            }
        }

        # ── Стилі ──
        self.setStyleSheet(GLOBAL_STYLESHEET)

        # ── Центральний стек екранів ──
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)

        self.stacked = QStackedWidget()
        self.main_layout.addWidget(self.stacked)

        # ── Створення екранів ──
        self.main_screen = self._build_main_screen()
        self.settings_screen = self._build_settings_screen()
        self.about_screen = self._build_about_screen()
        self.how_it_works_screen = self._build_info_screen(
            self.t("how_it_works_title"),
            self.t("how_it_works_text"),
            back_target="about"
        )
        self.features_screen = self._build_info_screen(
            self.t("features_title"),
            self.t("features_text"),
            back_target="about"
        )

        self.stacked.addWidget(self.main_screen)        # index 0
        self.stacked.addWidget(self.settings_screen)     # index 1
        self.stacked.addWidget(self.about_screen)        # index 2
        self.stacked.addWidget(self.how_it_works_screen) # index 3
        self.stacked.addWidget(self.features_screen)     # index 4

        self.stacked.setCurrentIndex(0)

        # ── Системний трей ──
        self._setup_tray()

    # ═══════════════════════════════════════════════════════════
    #                 SETTINGS (load / save)
    # ═══════════════════════════════════════════════════════════

    def _load_settings(self):
        default_path = os.path.join(os.path.expanduser("~"), "Desktop")
        defaults = {
            "save_path": default_path,
            "font_family": "Georgia",
            "font_size": 16,
            "output_format": "docx",
            "translation_engine": "Google Translator",
            "theme": "dark",
            "ui_language": "uk",
            "target_lang_name": "Українська",
            "add_read_time": True,
            "bilingual_mode": False,
            "auto_open": True,
            "download_images": True,
            "add_toc": True,
            "add_metadata": True,
            "minimize_to_tray": False
        }
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, "r", encoding="utf-8") as f:
                    loaded = json.load(f)
                    return {**defaults, **loaded}
            except Exception:
                return defaults
        return defaults

    def _save_settings(self):
        try:
            with open(self.config_file, "w", encoding="utf-8") as f:
                json.dump(self.state, f, ensure_ascii=False, indent=4)
        except Exception:
            pass

    # ═══════════════════════════════════════════════════════════
    #                       ХЕЛПЕРИ
    # ═══════════════════════════════════════════════════════════

    def t(self, key):
        lang = self.state.get("ui_language", "uk")
        return self.locales.get(lang, self.locales["uk"]).get(key, key)

    def navigate(self, screen_name):
        """Навігація між екранами"""
        mapping = {
            "main": 0,
            "settings": 1,
            "about": 2,
            "how_it_works": 3,
            "features": 4
        }
        idx = mapping.get(screen_name, 0)
        self.stacked.setCurrentIndex(idx)

    def is_valid_url(self, url):
        parsed = urllib.parse.urlparse(url)
        return all([parsed.scheme in ['http', 'https'], parsed.netloc])

    # ═══════════════════════════════════════════════════════════
    #               ЕКРАН 1: ГОЛОВНИЙ ЕКРАН
    # ═══════════════════════════════════════════════════════════

    def _build_main_screen(self):
        screen = QWidget()
        screen.setStyleSheet(f"background-color: {Theme.BG_PRIMARY};")
        layout = QVBoxLayout(screen)
        layout.setContentsMargins(40, 20, 40, 30)
        layout.setSpacing(0)

        # ── Верхня панель: кнопка налаштувань ──
        top_bar = QHBoxLayout()
        top_bar.addStretch()
        settings_btn = IconButton("⚙️", size=50, font_size=30)
        settings_btn.setToolTip(self.t("settings_title"))
        settings_btn.clicked.connect(lambda: self.navigate("settings"))
        top_bar.addWidget(settings_btn)
        layout.addLayout(top_bar)

        layout.addSpacing(15)

        # ── Іконка-логотип з пасхалкою ──
        temple_label = QLabel("🏛️")
        temple_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        temple_label.setFont(QFont("Arial", 42))
        temple_label.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        temple_label.mousePressEvent = self._on_temple_click
        layout.addWidget(temple_label)

        layout.addSpacing(5)

        # ── Заголовок з пасхалкою ──
        self.main_title_label = QLabel(self.t("title").split(" v")[0])
        self.main_title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.main_title_label.setFont(QFont(Theme.FONT_TITLE, 52, QFont.Weight.Bold))
        self.main_title_label.setStyleSheet(f"color: {Theme.GOLD};")
        self.main_title_label.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.main_title_label.mousePressEvent = self._on_title_click
        layout.addWidget(self.main_title_label)

        # ── Підзаголовок (версія) ──
        version_label = QLabel("v6.0")
        version_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        version_label.setFont(QFont(Theme.FONT_UI, 14))
        version_label.setStyleSheet(f"color: {Theme.TEXT_MUTED};")
        layout.addWidget(version_label)

        layout.addSpacing(25)

        # ── Поле вводу URL ──
        self.url_textbox = QTextEdit()
        self.url_textbox.setPlaceholderText(self.t("placeholder"))
        self.url_textbox.setFixedHeight(120)
        self.url_textbox.setFont(QFont(Theme.FONT_UI, 15))

        # Додаємо тінь до поля вводу
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(30)
        shadow.setColor(QColor(212, 175, 55, 40))
        shadow.setOffset(0, 4)
        self.url_textbox.setGraphicsEffect(shadow)
        layout.addWidget(self.url_textbox)

        layout.addSpacing(15)

        # ── Прогресбар ──
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedHeight(12)
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        layout.addSpacing(10)

        # ── Статус ──
        self.status_label = QLabel(self.t("status_wait"))
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setFont(QFont(Theme.FONT_UI, 16))
        self.status_label.setStyleSheet(f"color: {Theme.TEXT_SECONDARY}; font-style: italic;")
        layout.addWidget(self.status_label)

        layout.addSpacing(25)

        # ── Кнопки дій ──
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(16)

        self.save_btn = PrimaryButton(self.t("btn_digitize"))
        self.save_btn.setFixedWidth(380)
        self.save_btn.add_shadow()
        self.save_btn.clicked.connect(self._on_digitize_click)
        buttons_layout.addWidget(self.save_btn)

        self.cancel_btn = DangerButton(self.t("btn_cancel"))
        self.cancel_btn.setFixedWidth(200)
        self.cancel_btn.setEnabled(False)
        self.cancel_btn.clicked.connect(self._on_cancel_click)
        buttons_layout.addWidget(self.cancel_btn)

        buttons_container = QWidget()
        buttons_container.setLayout(buttons_layout)
        layout.addWidget(buttons_container, alignment=Qt.AlignmentFlag.AlignCenter)

        layout.addStretch()

        # ── Футер ──
        footer_label = QLabel("Made with ❤️ for free access to knowledge")
        footer_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        footer_label.setFont(QFont(Theme.FONT_UI, 11))
        footer_label.setStyleSheet(f"color: {Theme.TEXT_MUTED};")
        layout.addWidget(footer_label)

        return screen

    # ═══════════════════════════════════════════════════════════
    #              ЕКРАН 2: НАЛАШТУВАННЯ
    # ═══════════════════════════════════════════════════════════

    def _build_settings_screen(self):
        screen = QWidget()
        screen.setStyleSheet(f"background-color: {Theme.BG_PRIMARY};")
        outer_layout = QVBoxLayout(screen)
        outer_layout.setContentsMargins(40, 20, 40, 20)

        # ── Верхня панель ──
        top_bar = QHBoxLayout()
        back_btn = IconButton("←", size=50, font_size=34)
        back_btn.clicked.connect(lambda: self.navigate("main"))
        top_bar.addWidget(back_btn)
        top_bar.addStretch()

        title = QLabel(self.t("settings_title"))
        title.setFont(QFont(Theme.FONT_TITLE, 30, QFont.Weight.Bold))
        title.setStyleSheet(f"color: {Theme.GOLD};")
        top_bar.addWidget(title)
        top_bar.addStretch()
        top_bar.addSpacing(50)  # Баланс з кнопкою назад
        outer_layout.addLayout(top_bar)

        outer_layout.addSpacing(10)

        # ── Скролабельна область ──
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff
        )

        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(60, 10, 60, 30)
        scroll_layout.setSpacing(12)

        # ── Мова інтерфейсу ──
        scroll_layout.addWidget(SectionLabel(self.t("ui_lang_lbl")))
        self.ui_lang_combo = NoScrollComboBox()
        self.ui_lang_combo.addItems(["Українська", "English", "Ельфійська (Sindarin)"])
        current_lang = "Українська" if self.state["ui_language"] == "uk" else "English"
        self.ui_lang_combo.setCurrentText(current_lang)
        self.ui_lang_combo.currentTextChanged.connect(self._on_ui_lang_changed)
        scroll_layout.addWidget(self.ui_lang_combo)

        scroll_layout.addSpacing(8)

        # ── Мова перекладу ──
        scroll_layout.addWidget(SectionLabel(self.t("target_lang_lbl")))
        self.target_combo = NoScrollComboBox()
        self.target_combo.addItems(["Українська", "English", "Polski", "Deutsch", "Français", "Español"])
        self.target_combo.setCurrentText(self.state.get("target_lang_name", "Українська"))
        self.target_combo.currentTextChanged.connect(
            lambda v: self._update_state("target_lang_name", v)
        )
        scroll_layout.addWidget(self.target_combo)

        scroll_layout.addSpacing(8)

        # ── Рушій перекладу ──
        scroll_layout.addWidget(SectionLabel(self.t("engine_lbl")))
        self.engine_combo = NoScrollComboBox()
        self.engine_combo.addItems(["Google Translator", "Microsoft Translator", "Skynet v2.0"])
        self.engine_combo.setCurrentText(self.state.get("translation_engine", "Google Translator"))
        self.engine_combo.currentTextChanged.connect(self._on_engine_changed)
        scroll_layout.addWidget(self.engine_combo)

        scroll_layout.addSpacing(8)

        # ── Папка збереження ──
        scroll_layout.addWidget(SectionLabel(self.t("path_lbl")))
        path_row = QHBoxLayout()
        self.path_display = QLabel(self.state["save_path"])
        self.path_display.setFont(QFont(Theme.FONT_UI, 13))
        self.path_display.setStyleSheet(f"""
            color: {Theme.TEXT_SECONDARY};
            background-color: {Theme.BG_INPUT};
            border: 1px solid {Theme.BORDER};
            border-radius: 10px;
            padding: 10px 14px;
        """)
        self.path_display.setWordWrap(True)
        path_row.addWidget(self.path_display, stretch=1)

        browse_btn = GoldButton(self.t("btn_choose"))
        browse_btn.setFixedWidth(150)
        browse_btn.setFixedHeight(42)
        browse_btn.clicked.connect(self._on_browse_folder)
        path_row.addWidget(browse_btn)
        scroll_layout.addLayout(path_row)

        scroll_layout.addSpacing(8)

        # ── Шрифт ──
        scroll_layout.addWidget(SectionLabel(self.t("font_lbl")))
        self.font_combo = NoScrollComboBox()
        self.font_combo.addItems(["Georgia", "Arial", "Times New Roman", "Comic Sans MS"])
        self.font_combo.setCurrentText(self.state.get("font_family", "Georgia"))
        self.font_combo.currentTextChanged.connect(self._on_font_changed)
        scroll_layout.addWidget(self.font_combo)

        scroll_layout.addSpacing(8)

        # ── Формат ──
        scroll_layout.addWidget(SectionLabel(self.t("format_lbl")))
        self.format_combo = NoScrollComboBox()
        self.format_combo.addItems(["docx", "pdf"])
        self.format_combo.setCurrentText(self.state.get("output_format", "docx"))
        self.format_combo.currentTextChanged.connect(self._on_format_changed)
        scroll_layout.addWidget(self.format_combo)

        scroll_layout.addSpacing(8)

        # ── Розмір тексту (слайдер) ──
        size_row = QHBoxLayout()
        size_row.addWidget(SectionLabel(self.t("size_lbl")))
        size_row.addStretch()
        self.size_value_label = QLabel(str(self.state["font_size"]))
        self.size_value_label.setFont(QFont(Theme.FONT_UI, 18, QFont.Weight.Bold))
        self.size_value_label.setStyleSheet(f"color: {Theme.GOLD};")
        size_row.addWidget(self.size_value_label)
        scroll_layout.addLayout(size_row)

        size_slider = NoScrollSlider(Qt.Orientation.Horizontal)
        size_slider.setRange(12, 24)
        size_slider.setValue(self.state["font_size"])
        size_slider.setTickInterval(1)
        size_slider.valueChanged.connect(self._on_font_size_changed)
        scroll_layout.addWidget(size_slider)

        scroll_layout.addSpacing(16)

        # ── Додаткові функції (Toggle Switches) ──
        scroll_layout.addWidget(GoldSectionLabel(self.t("additional_features_lbl")))

        self.toggle_read_time = ToggleSwitch(self.t("setting_read_time"), self.state["add_read_time"])
        self.toggle_read_time.toggled.connect(lambda v: self._update_state("add_read_time", v))
        scroll_layout.addWidget(self.toggle_read_time)

        self.toggle_toc = ToggleSwitch(self.t("setting_toc"), self.state["add_toc"])
        self.toggle_toc.toggled.connect(lambda v: self._update_state("add_toc", v))
        scroll_layout.addWidget(self.toggle_toc)

        self.toggle_metadata = ToggleSwitch(self.t("setting_metadata"), self.state["add_metadata"])
        self.toggle_metadata.toggled.connect(lambda v: self._update_state("add_metadata", v))
        scroll_layout.addWidget(self.toggle_metadata)

        self.toggle_bilingual = ToggleSwitch(self.t("setting_bilingual"), self.state["bilingual_mode"])
        self.toggle_bilingual.toggled.connect(self._on_bilingual_toggled)
        scroll_layout.addWidget(self.toggle_bilingual)

        self.toggle_images = ToggleSwitch(self.t("setting_images"), self.state["download_images"])
        self.toggle_images.toggled.connect(lambda v: self._update_state("download_images", v))
        scroll_layout.addWidget(self.toggle_images)

        self.toggle_auto_open = ToggleSwitch(self.t("setting_auto_open"), self.state["auto_open"])
        self.toggle_auto_open.toggled.connect(lambda v: self._update_state("auto_open", v))
        scroll_layout.addWidget(self.toggle_auto_open)

        self.toggle_tray = ToggleSwitch(self.t("setting_tray_close"), self.state["minimize_to_tray"])
        self.toggle_tray.toggled.connect(lambda v: self._update_state("minimize_to_tray", v))
        scroll_layout.addWidget(self.toggle_tray)

        scroll_layout.addSpacing(20)

        # ── Тема ──
        scroll_layout.addWidget(SectionLabel(self.t("theme_lbl")))
        self.toggle_theme = ToggleSwitch("🌙 " + self.t("theme_dark"), self.state["theme"] == "dark")
        self.toggle_theme.toggled.connect(self._on_theme_toggled)
        scroll_layout.addWidget(self.toggle_theme)

        scroll_layout.addSpacing(30)

        about_btn = GoldButton(self.t("btn_about"))
        about_btn.setMinimumWidth(250)
        about_btn.clicked.connect(lambda: self.navigate("about"))
        scroll_layout.addWidget(about_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        scroll_layout.addSpacing(20)

        # ── Premium Easter Egg ──
        premium_btn = QPushButton("v6.0 👑")
        premium_btn.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        premium_btn.setFont(QFont(Theme.FONT_UI, 11))
        premium_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                color: {Theme.TEXT_MUTED};
                border: none;
            }}
            QPushButton:hover {{
                color: {Theme.GOLD};
            }}
        """)
        premium_btn.clicked.connect(self._show_premium_joke)
        scroll_layout.addWidget(premium_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        scroll_layout.addSpacing(20)

        scroll_area.setWidget(scroll_content)
        outer_layout.addWidget(scroll_area)

        return screen

    # ═══════════════════════════════════════════════════════════
    #              ЕКРАН 3: ПРО ДОДАТОК
    # ═══════════════════════════════════════════════════════════

    def _build_about_screen(self):
        screen = QWidget()
        screen.setStyleSheet(f"background-color: {Theme.BG_PRIMARY};")
        layout = QVBoxLayout(screen)
        layout.setContentsMargins(50, 20, 50, 30)

        # ── Верхня панель ──
        top_bar = QHBoxLayout()
        back_btn = IconButton("←", size=50, font_size=34)
        back_btn.clicked.connect(lambda: self.navigate("settings"))
        top_bar.addWidget(back_btn)
        top_bar.addStretch()
        outer_title = QLabel(self.t("about_title"))
        outer_title.setFont(QFont(Theme.FONT_TITLE, 30, QFont.Weight.Bold))
        outer_title.setStyleSheet(f"color: {Theme.GOLD};")
        top_bar.addWidget(outer_title)
        top_bar.addStretch()
        top_bar.addSpacing(50)
        layout.addLayout(top_bar)

        layout.addSpacing(15)

        # ── Іконка ──
        icon_lbl = QLabel("🏛️")
        icon_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        icon_lbl.setFont(QFont("Arial", 55))
        layout.addWidget(icon_lbl)

        layout.addSpacing(10)

        # ── Опис ──
        desc_label = QLabel(self.t("about_desc"))
        desc_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        desc_label.setFont(QFont(Theme.FONT_UI, 14))
        desc_label.setStyleSheet(f"color: {Theme.TEXT_SECONDARY}; line-height: 1.6;")
        desc_label.setWordWrap(True)
        desc_label.setMaximumWidth(750)
        desc_label.setMinimumHeight(150)
        layout.addWidget(desc_label, alignment=Qt.AlignmentFlag.AlignCenter)

        layout.addSpacing(30)

        # ── Перший ряд кнопок ──
        row1 = QHBoxLayout()
        row1.setSpacing(12)

        btn_how = FeatureButton(
            self.t("btn_how_it_works"),
            Theme.BLUE, Theme.BLUE_HOVER
        )
        btn_how.clicked.connect(lambda: self.navigate("how_it_works"))
        row1.addWidget(btn_how)

        btn_features = FeatureButton(
            self.t("btn_features"),
            Theme.GREEN, Theme.GREEN_HOVER
        )
        btn_features.clicked.connect(lambda: self.navigate("features"))
        row1.addWidget(btn_features)

        row1_wrapper = QWidget()
        row1_wrapper.setLayout(row1)
        layout.addWidget(row1_wrapper, alignment=Qt.AlignmentFlag.AlignCenter)

        layout.addSpacing(10)

        # ── Другий ряд кнопок ──
        row2 = QHBoxLayout()
        row2.setSpacing(12)

        btn_donate = FeatureButton(
            self.t("btn_donate"),
            Theme.ORANGE, Theme.ORANGE_HOVER
        )
        btn_donate.clicked.connect(self._on_donate_click)
        row2.addWidget(btn_donate)

        btn_feedback = FeatureButton(
            self.t("btn_feedback"),
            "#4b4b4b", Theme.BG_HOVER
        )
        btn_feedback.clicked.connect(self._on_feedback_click)
        row2.addWidget(btn_feedback)

        row2_wrapper = QWidget()
        row2_wrapper.setLayout(row2)
        layout.addWidget(row2_wrapper, alignment=Qt.AlignmentFlag.AlignCenter)

        layout.addStretch()

        return screen

    # ═══════════════════════════════════════════════════════════
    #          ЕКРАНИ 4-5: ІНФОРМАЦІЙНІ (How it works / Features)
    # ═══════════════════════════════════════════════════════════

    def _build_info_screen(self, title_text, body_text, back_target="about"):
        screen = QWidget()
        screen.setStyleSheet(f"background-color: {Theme.BG_PRIMARY};")
        layout = QVBoxLayout(screen)
        layout.setContentsMargins(40, 20, 40, 30)

        # ── Верхня панель ──
        top_bar = QHBoxLayout()
        back_btn = IconButton("←", size=50, font_size=34)
        back_btn.clicked.connect(lambda: self.navigate(back_target))
        top_bar.addWidget(back_btn)
        top_bar.addStretch()
        title_lbl = QLabel(title_text)
        title_lbl.setFont(QFont(Theme.FONT_TITLE, 28, QFont.Weight.Bold))
        title_lbl.setStyleSheet(f"color: {Theme.GOLD};")
        top_bar.addWidget(title_lbl)
        top_bar.addStretch()
        top_bar.addSpacing(50)
        layout.addLayout(top_bar)

        layout.addSpacing(30)

        # ── Контент ──
        content_card = CardFrame()
        card_layout = QVBoxLayout(content_card)
        card_layout.setContentsMargins(30, 25, 30, 25)

        text_label = QLabel(body_text)
        text_label.setFont(QFont(Theme.FONT_UI, 15))
        text_label.setStyleSheet(f"color: {Theme.TEXT_PRIMARY}; line-height: 1.7;")
        text_label.setWordWrap(True)
        card_layout.addWidget(text_label)

        layout.addWidget(content_card)
        layout.addStretch()

        return screen

    # ═══════════════════════════════════════════════════════════
    #                   СИСТЕМНИЙ ТРЕЙ
    # ═══════════════════════════════════════════════════════════

    def _setup_tray(self):
        """Налаштування іконки в системному треї"""
        self.tray_icon = QSystemTrayIcon(self)

        # Створюємо простий піксмап якщо іконки немає
        try:
            icon = QIcon("icon.ico")
            if icon.isNull():
                raise FileNotFoundError
        except Exception:
            pixmap = QPixmap(64, 64)
            pixmap.fill(QColor(Theme.BG_CARD))
            painter = QPainter(pixmap)
            painter.setBrush(QBrush(QColor(Theme.GOLD)))
            painter.setPen(Qt.PenStyle.NoPen)
            painter.drawRoundedRect(8, 8, 48, 48, 8, 8)
            painter.end()
            icon = QIcon(pixmap)

        self.tray_icon.setIcon(icon)
        self.tray_icon.setToolTip(self.t("title"))

        # Меню трею
        tray_menu = QMenu()
        tray_menu.setStyleSheet(f"""
            QMenu {{
                background-color: {Theme.BG_CARD};
                color: {Theme.TEXT_PRIMARY};
                border: 1px solid {Theme.BORDER};
                border-radius: 8px;
                padding: 6px;
            }}
            QMenu::item {{
                padding: 8px 20px;
                border-radius: 4px;
            }}
            QMenu::item:selected {{
                background-color: {Theme.GOLD_DARK};
            }}
        """)

        action_open = tray_menu.addAction(self.t("qm_btn_open"))
        action_open.triggered.connect(self._show_from_tray)

        action_clipboard = tray_menu.addAction(self.t("qm_btn_paste"))
        action_clipboard.triggered.connect(self._tray_clipboard_action)

        tray_menu.addSeparator()

        action_quit = tray_menu.addAction(self.t("qm_btn_exit"))
        action_quit.triggered.connect(self._quit_app)

        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.activated.connect(self._on_tray_activated)
        self.tray_icon.show()

    def _show_from_tray(self):
        self.showNormal()
        self.activateWindow()

    def _on_tray_activated(self, reason):
        if reason == QSystemTrayIcon.ActivationReason.DoubleClick:
            self._show_from_tray()

    def _tray_clipboard_action(self):
        """Обробка статті з буфера обміну через трей"""
        if self.is_processing:
            self.tray_icon.showMessage(
                "Скарбниця Знань",
                self.t("tray_processing_err"),
                QSystemTrayIcon.MessageIcon.Warning,
                3000
            )
            return

        clipboard_text = ""
        if pyperclip:
            try:
                clipboard_text = pyperclip.paste().strip()
            except Exception:
                pass

        if not clipboard_text:
            try:
                clipboard = QApplication.clipboard()
                clipboard_text = clipboard.text().strip()
            except Exception:
                pass

        # ── Пасхалка: пароль ──
        if clipboard_text and any(w in clipboard_text.lower() for w in ["пароль", "password", "123456"]):
            self.tray_icon.showMessage(
                "Скарбниця Знань",
                "Я парсер статей, а не крадій паролів! (Але я його запам'ятав 🕵️‍♂️)",
                QSystemTrayIcon.MessageIcon.Information,
                3000
            )
            return

        if not self.is_valid_url(clipboard_text):
            self.tray_icon.showMessage(
                "Скарбниця Знань",
                self.t("tray_clipboard_err"),
                QSystemTrayIcon.MessageIcon.Warning,
                3000
            )
            return

        self.tray_icon.showMessage(
            "Скарбниця Знань",
            f"Фонова обробка: {clipboard_text[:40]}...",
            QSystemTrayIcon.MessageIcon.Information,
            3000
        )

        self._start_worker([clipboard_text])

    def _quit_app(self):
        self.tray_icon.hide()
        QApplication.instance().quit()

    # ═══════════════════════════════════════════════════════════
    #            SETTINGS CHANGE HANDLERS
    # ═══════════════════════════════════════════════════════════

    def _update_state(self, key, value):
        """Універсальний метод для оновлення стану та збереження"""
        self.state[key] = value
        self._save_settings()

    def _on_ui_lang_changed(self, choice):
        if choice == "Ельфійська (Sindarin)":
            QMessageBox.critical(
                self, "Мордорська помилка",
                "Помилка API: Словник ельфійської мови згорів у Мордорі. "
                "Будь ласка, оберіть людську мову 🧝‍♂️🔥"
            )
            self.ui_lang_combo.setCurrentText(
                "Українська" if self.state["ui_language"] == "uk" else "English"
            )
            return

        new_lang = "uk" if choice == "Українська" else "en"
        self.state["ui_language"] = new_lang
        self._save_settings()
        self.setWindowTitle(self.t("title"))
        # Перебудовуємо UI для нової мови
        self._rebuild_all_screens()

    def _on_engine_changed(self, choice):
        if choice == "Skynet v2.0":
            QMessageBox.warning(
                self, "Skynet",
                "Помилка: Skynet зараз зайнятий захопленням світу. "
                "Повертаємось до Google 🤖🔥"
            )
            self.engine_combo.setCurrentText("Google Translator")
            return
        self._update_state("translation_engine", choice)

    def _on_font_changed(self, choice):
        if choice == "Comic Sans MS":
            QMessageBox.information(
                self, "Шрифт Лікарів",
                "Серйозно? Comic Sans? Ваші очі вам цього не пробачать... але як скажете 🤡"
            )
        self._update_state("font_family", choice)

    def _on_format_changed(self, choice):
        curr_time = time.time()
        if curr_time - self.format_click_time < 2.0:
            self.format_clicks += 1
            if self.format_clicks >= 4:
                QMessageBox.information(
                    self, "Криза ідентичності",
                    "Визначся вже! Я тобі Microsoft Word чи Acrobat Reader? 🤯📄"
                )
                self.format_clicks = 0
        else:
            self.format_clicks = 1
        self.format_click_time = curr_time
        self._update_state("output_format", choice)

    def _on_bilingual_toggled(self, checked):
        curr_time = time.time()
        if curr_time - self.bilingual_click_time < 1.0:
            self.bilingual_clicks += 1
            if self.bilingual_clicks >= 4:
                self.toggle_bilingual.label.setText(
                    "My brain is melting. Мій мозок плавиться. 🧠🔥"
                )
                self.toggle_bilingual.label.setStyleSheet(f"color: {Theme.RED};")
                self.bilingual_clicks = 0
        else:
            self.bilingual_clicks = 1
        self.bilingual_click_time = curr_time
        self._update_state("bilingual_mode", checked)

    def _on_theme_toggled(self, checked):
        curr_time = time.time()
        if curr_time - self.theme_click_time < 1.0:
            self.theme_clicks += 1
            if self.theme_clicks >= 5:
                QMessageBox.warning(
                    self, "Світломузика",
                    "Зупинись! Ти хочеш, щоб у мене стався епілептичний напад? "
                    "Залишаємо темну! 😵‍💫🕶️"
                )
                self.toggle_theme.setChecked(True)
                self.state["theme"] = "dark"
                self._save_settings()
                self.theme_clicks = 0
                return
        else:
            self.theme_clicks = 1
        self.theme_click_time = curr_time

        new_theme = "dark" if checked else "light"
        self._update_state("theme", new_theme)

    def _on_font_size_changed(self, value):
        """Зміна розміру шрифту"""
        self.state["font_size"] = value
        self._save_settings()
        if value == 24:
            self.font_slider_clicks += 1
            if self.font_slider_clicks >= 3:
                self.size_value_label.setText(
                    "24 (Ти збираєшся читати це з іншої кімнати? 🔭)"
                )
                return
        else:
            self.font_slider_clicks = 0
        self.size_value_label.setText(str(value))

    def _on_browse_folder(self):
        """Вибір папки збереження"""
        folder = QFileDialog.getExistingDirectory(self, "Оберіть папку")
        if folder:
            self._update_state("save_path", folder)
            self.path_display.setText(folder)
            self.cancel_folder_clicks = 0
        else:
            self.cancel_folder_clicks += 1
            if self.cancel_folder_clicks >= 3:
                QMessageBox.information(
                    self, "Шлях у нікуди",
                    "Які ми сьогодні нерішучі... "
                    "Зберігай на Робочий стіл, як всі нормальні люди! 🤷‍♂️"
                )
                self.cancel_folder_clicks = 0

    def _rebuild_all_screens(self):
        """Перебудовує всі екрани після зміни мови"""
        current_idx = self.stacked.currentIndex()

        # Видаляємо старі екрани
        while self.stacked.count() > 0:
            widget = self.stacked.widget(0)
            self.stacked.removeWidget(widget)
            widget.deleteLater()

        # Будуємо заново
        self.main_screen = self._build_main_screen()
        self.settings_screen = self._build_settings_screen()
        self.about_screen = self._build_about_screen()
        self.how_it_works_screen = self._build_info_screen(
            self.t("how_it_works_title"),
            self.t("how_it_works_text"),
            back_target="about"
        )
        self.features_screen = self._build_info_screen(
            self.t("features_title"),
            self.t("features_text"),
            back_target="about"
        )

        self.stacked.addWidget(self.main_screen)
        self.stacked.addWidget(self.settings_screen)
        self.stacked.addWidget(self.about_screen)
        self.stacked.addWidget(self.how_it_works_screen)
        self.stacked.addWidget(self.features_screen)

        self.stacked.setCurrentIndex(current_idx)

    # ═══════════════════════════════════════════════════════════
    #            ОСНОВНА ЛОГІКА (DIGITIZE / CANCEL)
    # ═══════════════════════════════════════════════════════════

    def _on_digitize_click(self):
        """Натискання кнопки ОЦИФРУВАТИ"""
        url_text = self.url_textbox.toPlainText().strip()

        if not url_text:
            self.empty_clicks += 1
            if self.empty_clicks == 3:
                self._set_status("Я все ще чекаю посилання...", Theme.GOLD)
            elif self.empty_clicks >= 5:
                self._set_status(
                    "Слухай, я парсер, а не телепат. Дай лінк! 🤬",
                    Theme.RED
                )
            return

        self.empty_clicks = 0

        # Пасхалки на ввід
        if url_text.lower() in ["wake up", "matrix"]:
            self._trigger_matrix_effect()
            return

        if "youtube.com/watch?v=dQw4w9WgXcQ" in url_text or url_text.lower() == "кава":
            self._set_status(
                "Гарна спроба! Але мене не зарікролити 😎 "
                "(або ти просто хочеш пригостити мене кавою)",
                Theme.GOLD
            )
            if url_text.lower() == "кава":
                self._on_donate_click()
            return

        urls = [u.strip() for u in url_text.split('\n') if u.strip()]
        valid_urls = [u for u in urls if self.is_valid_url(u)]
        if not valid_urls:
            QMessageBox.warning(self, "Увага", self.t("msg_invalid_url"))
            return

        self.start_process_time = time.time()
        self._start_worker(valid_urls)

    def _start_worker(self, urls):
        """Запускає ScrapingWorker у фоновому потоці"""
        self.is_processing = True
        self.save_btn.setEnabled(False)
        self.cancel_btn.setEnabled(True)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self._set_status(self.t("status_single_start"), Theme.GOLD)

        self.worker = ScrapingWorker(urls, self.state, self.locales)

        # Підключаємо сигнали
        self.worker.progress_updated.connect(self._on_worker_progress)
        self.worker.status_updated.connect(self._on_worker_status)
        self.worker.finished_all.connect(self._on_worker_finished)
        self.worker.error_occurred.connect(self._on_worker_error)
        self.worker.notification_requested.connect(self._on_worker_notification)

        self.worker.start()

    def _on_worker_progress(self, value):
        self.progress_bar.setValue(value)

    def _on_worker_status(self, text, color):
        self._set_status(text, color)

    def _on_worker_finished(self):
        self.is_processing = False
        self.save_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        QTimer.singleShot(3000, lambda: self.progress_bar.setVisible(False))

    def _on_worker_error(self, error_msg):
        QMessageBox.critical(self, "Error", error_msg)

    def _on_worker_notification(self, title, message):
        self.tray_icon.showMessage(
            title, message,
            QSystemTrayIcon.MessageIcon.Information,
            5000
        )

    def _on_cancel_click(self):
        """Скасування"""
        if self.worker and self.worker.isRunning():
            self.worker.cancel()

        elapsed = time.time() - self.start_process_time
        if elapsed < 2.0:
            self._set_status(
                "🛑 Скасовано. Ти що, забув вимкнути праску вдома? 🏃‍♂️💨",
                Theme.RED
            )
        else:
            self._set_status(self.t("status_cancelled"), Theme.RED)

        self.is_processing = False
        self.save_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.progress_bar.setVisible(False)

    # ═══════════════════════════════════════════════════════════
    #                  ABOUT / DONATE / FEEDBACK
    # ═══════════════════════════════════════════════════════════

    def _on_donate_click(self):
        webbrowser.open("https://send.monobank.ua/jar/328DrBEZXY")

    def _on_feedback_click(self):
        email = "treasuryofknowledge26@gmail.com"
        subject = "Відгук про Скарбницю Знань v6.0"

        # Копіюємо в буфер
        if pyperclip:
            try:
                pyperclip.copy(email)
            except Exception:
                pass

        try:
            webbrowser.open(
                f"mailto:{email}?subject={urllib.parse.quote(subject)}", new=1
            )
        except Exception:
            pass

        QMessageBox.information(
            self, "Зворотний зв'язок",
            f"Адресу {email} скопійовано в буфер обміну!\n\n"
            "Якщо поштова програма не відкрилася, "
            "вставте адресу вручну в поле 'Кому'."
        )

    # ═══════════════════════════════════════════════════════════
    #                    ПАСХАЛКИ
    # ═══════════════════════════════════════════════════════════

    def _on_temple_click(self, event):
        self.temple_clicks += 1
        if self.temple_clicks == 7:
            self.hacker_mode = True
            self._set_status(
                "Ви знайшли секретний підвал Храму! "
                "Тепер ви Верховний Жрець Знань 🧙‍♂️",
                "#00FF00"
            )
            self.temple_clicks = 0

    def _on_title_click(self, event):
        self.title_clicks += 1
        if self.title_clicks == 3:
            self.main_title_label.setText("Бібліотека Піратів 🏴‍☠️")
            QTimer.singleShot(
                2000,
                lambda: self.main_title_label.setText("Робін Гуд PDF-формату 🏹")
            )
            QTimer.singleShot(
                4000,
                lambda: self.main_title_label.setText(
                    self.t("title").split(" v")[0]
                )
            )
            self.title_clicks = 0

    def _trigger_matrix_effect(self):
        """The Matrix has you..."""
        self.url_textbox.clear()
        msg = "The Matrix has you... Follow the white rabbit 🐇"
        self._matrix_idx = 0

        def type_next():
            if self._matrix_idx < len(msg):
                self.url_textbox.insertPlainText(msg[self._matrix_idx])
                self._matrix_idx += 1
                QTimer.singleShot(80, type_next)
            else:
                QTimer.singleShot(
                    3000,
                    lambda: self.url_textbox.clear()
                )

        type_next()

    def _show_premium_joke(self):
        """Easter Egg: Premium"""
        QMessageBox.information(
            self,
            self.t("premium_title"),
            f"💎\n\n{self.t('premium_text')}"
        )

    # ═══════════════════════════════════════════════════════════
    #                    УТИЛІТИ
    # ═══════════════════════════════════════════════════════════

    def _set_status(self, text, color=None):
        """Оновлення статусу на головному екрані"""
        self.status_label.setText(text)
        if color:
            self.status_label.setStyleSheet(
                f"color: {color}; font-style: italic; font-size: 16px;"
            )
        else:
            self.status_label.setStyleSheet(
                f"color: {Theme.TEXT_SECONDARY}; font-style: italic; font-size: 16px;"
            )

    def closeEvent(self, event):
        """Обробка закриття вікна"""
        if self.state.get("minimize_to_tray", False):
            event.ignore()
            self.hide()
            self.tray_icon.showMessage(
                "Скарбниця Знань",
                "Програму згорнуто в трей. Готова до фонової роботи.",
                QSystemTrayIcon.MessageIcon.Information,
                3000
            )
        else:
            # Зупиняємо воркер, якщо він працює
            if self.worker and self.worker.isRunning():
                self.worker.cancel()
                self.worker.wait(3000)
            self.tray_icon.hide()
            event.accept()


# ═══════════════════════════════════════════════════════════════
#                        ЗАПУСК
# ═══════════════════════════════════════════════════════════════

def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    # Темна палітра на рівні Qt
    palette = QPalette()
    palette.setColor(QPalette.ColorRole.Window, QColor(Theme.BG_PRIMARY))
    palette.setColor(QPalette.ColorRole.WindowText, QColor(Theme.TEXT_PRIMARY))
    palette.setColor(QPalette.ColorRole.Base, QColor(Theme.BG_INPUT))
    palette.setColor(QPalette.ColorRole.AlternateBase, QColor(Theme.BG_SECONDARY))
    palette.setColor(QPalette.ColorRole.ToolTipBase, QColor(Theme.BG_CARD))
    palette.setColor(QPalette.ColorRole.ToolTipText, QColor(Theme.GOLD))
    palette.setColor(QPalette.ColorRole.Text, QColor(Theme.TEXT_PRIMARY))
    palette.setColor(QPalette.ColorRole.Button, QColor(Theme.BG_CARD))
    palette.setColor(QPalette.ColorRole.ButtonText, QColor(Theme.TEXT_PRIMARY))
    palette.setColor(QPalette.ColorRole.Highlight, QColor(Theme.GOLD_DARK))
    palette.setColor(QPalette.ColorRole.HighlightedText, QColor("white"))
    app.setPalette(palette)

    window = TreasuryApp()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()